"""
file_extractor.py — Extraction de texte depuis des fichiers PDF, Word et images.

Utilisé par l'endpoint /api/v1/extract-text pour transformer un document
importé depuis le dashboard en texte brut, avant de l'envoyer dans la
pipeline d'analyse CNSA habituelle.

Formats supportés :
    .pdf            — via pypdf
    .doc / .docx    — via python-docx
    .jpg/.jpeg/.png — via pytesseract (OCR), avec fallback PIL si non disponible
"""

import io
import logging
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

EXTENSIONS_PDF    = {".pdf"}
EXTENSIONS_DOCX   = {".doc", ".docx"}
EXTENSIONS_IMAGE  = {".jpg", ".jpeg", ".png"}
EXTENSIONS_AUTORISES = EXTENSIONS_PDF | EXTENSIONS_DOCX | EXTENSIONS_IMAGE


def extraire_texte(nom_fichier: str, contenu: bytes) -> dict[str, Any]:
    """
    Extrait le texte brut d'un fichier PDF, Word ou image.

    Args:
        nom_fichier : Nom original du fichier (avec extension).
        contenu     : Contenu binaire du fichier.

    Returns:
        Dictionnaire avec :
            texte         — texte extrait (str)
            nb_pages      — nombre de pages/sections (int)
            avertissement — message si extraction partielle ou vide (str | None)
    """
    ext = Path(nom_fichier).suffix.lower()

    if ext in EXTENSIONS_PDF:
        return _extraire_pdf(contenu)
    elif ext in EXTENSIONS_DOCX:
        return _extraire_docx(contenu)
    elif ext in EXTENSIONS_IMAGE:
        return _extraire_image(contenu, ext)
    else:
        exts_str = "PDF (.pdf), Word (.docx), image (.jpg, .jpeg, .png)"
        raise ValueError(f"Format non supporté : '{ext}'. Formats acceptés : {exts_str}.")


def _extraire_pdf(contenu: bytes) -> dict[str, Any]:
    """Extraction texte depuis un PDF via pypdf."""
    from pypdf import PdfReader

    reader   = PdfReader(io.BytesIO(contenu))
    nb_pages = len(reader.pages)
    parties  = []

    for i, page in enumerate(reader.pages):
        try:
            texte_page = page.extract_text() or ""
            if texte_page.strip():
                parties.append(texte_page)
        except Exception as e:
            logger.warning(f"Erreur extraction page {i+1} PDF : {e}")

    texte = "\n\n".join(parties).strip()

    avertissement = None
    if not texte:
        avertissement = (
            "Aucun texte extrait — ce PDF semble être un scan (image). "
            "Veuillez saisir les informations manuellement dans l'onglet 'Saisie libre'."
        )
    elif len(texte) < 100:
        avertissement = (
            "Le texte extrait est très court. Vérifiez que le PDF contient bien du texte sélectionnable."
        )

    logger.info(f"PDF extrait | {nb_pages} pages | {len(texte)} caractères")
    return {"texte": texte, "nb_pages": nb_pages, "avertissement": avertissement}


def _extraire_docx(contenu: bytes) -> dict[str, Any]:
    """Extraction texte depuis un fichier Word (.docx) via python-docx."""
    import docx

    doc      = docx.Document(io.BytesIO(contenu))
    parties  = []

    # Paragraphes normaux
    for para in doc.paragraphs:
        if para.text.strip():
            parties.append(para.text.strip())

    # Texte dans les tableaux
    for table in doc.tables:
        for row in table.rows:
            cellules = [c.text.strip() for c in row.cells if c.text.strip()]
            if cellules:
                parties.append(" | ".join(cellules))

    texte    = "\n\n".join(parties).strip()
    nb_pages = len(doc.paragraphs)   # approximatif pour Word

    avertissement = None
    if not texte:
        avertissement = (
            "Le fichier Word ne contient pas de texte lisible. "
            "Veuillez saisir les informations manuellement."
        )

    logger.info(f"DOCX extrait | {len(doc.paragraphs)} paragraphes | {len(texte)} caractères")
    return {"texte": texte, "nb_pages": nb_pages, "avertissement": avertissement}


def _extraire_image(contenu: bytes, ext: str) -> dict[str, Any]:
    """OCR sur une image via pytesseract (si installé) ou fallback PIL."""
    texte = ""
    avertissement = None

    try:
        import pytesseract
        from PIL import Image
        img   = Image.open(io.BytesIO(contenu))
        texte = pytesseract.image_to_string(img, lang="fra+eng").strip()
        logger.info(f"Image OCR (tesseract) | ext={ext} | {len(texte)} caractères")
    except ImportError:
        # pytesseract non installé — tenter extraction PIL basique
        try:
            from PIL import Image
            img = Image.open(io.BytesIO(contenu))
            texte = ""
            logger.info(f"Image reçue (PIL uniquement, pas d'OCR) | ext={ext}")
        except ImportError:
            pass
    except Exception as e:
        logger.warning(f"Erreur OCR image : {e}")

    if not texte:
        avertissement = (
            "Aucun texte extrait de cette image. "
            "pytesseract n'est pas installé ou l'image ne contient pas de texte lisible. "
            "Veuillez saisir les informations manuellement dans l'onglet 'Saisie libre'."
        )

    return {"texte": texte, "nb_pages": 1, "avertissement": avertissement}
