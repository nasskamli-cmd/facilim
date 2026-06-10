"""
app/main.py — Point d'entrée FastAPI Facilim v2.

Architecture event-driven :
  - Webhook WhatsApp → OrchestrationEngine
  - Dashboard ESSMS  → API JSON sécurisée JWT
  - Portail usager   → (à venir) HDS

Tous les événements sont :
  - immédiatement persistés
  - journalisés dans l'audit trail
  - traçables end-to-end

Ce fichier ne contient QUE le câblage HTTP.
La logique métier est dans les engines/.
"""

from __future__ import annotations

import json
import logging
import threading
import time
import uuid
from datetime import datetime, timezone
from typing import Any

import time
import random
import threading

from fastapi import (
    BackgroundTasks, Depends, FastAPI, HTTPException, Request, Response,
    UploadFile, File, Form,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from app.config.settings import get_settings
from app.database import connection as db_conn_module
from app.security import encryption, token_manager, access_control
from app.audit import event_logger
from app.services.whatsapp_service import WhatsAppService, parse_webhook_payload
from app.services.notification_service import NotificationService
from app.engines.orchestration_engine import OrchestrationEngine

settings = get_settings()

# ── Store 2FA : codes de vérification en mémoire ─────────────────────────────
# {email: {"code": "123456", "expires": float, "jwt_token": str}}
_2FA_STORE:     dict[str, dict] = {}
_2FA_LOCK:      threading.Lock  = threading.Lock()
_2FA_TTL_SEC:   int             = 600   # 10 minutes


def _send_email_brevo(
    to_email: str,
    subject: str,
    html: str,
    reply_to_email: str | None = None,
    reply_to_name: str | None = None,
) -> bool:
    """
    Envoi d'email générique via l'API HTTP Brevo (réutilisable).
    Retourne True si Brevo a accepté l'envoi. Non bloquant côté appelant.
    """
    if not settings.brevo_api_key:
        logger.warning("[EMAIL] BREVO_API_KEY absente — envoi ignoré (to=%s)", (to_email or "")[:4] + "***")
        return False
    try:
        import requests as _req
        _payload = {
            "sender":      {"name": settings.brevo_sender_name, "email": settings.brevo_sender_email},
            "to":          [{"email": to_email}],
            "subject":     subject,
            "htmlContent": html,
        }
        if reply_to_email:
            _payload["replyTo"] = {"email": reply_to_email, "name": reply_to_name or reply_to_email}
        resp = _req.post(
            "https://api.brevo.com/v3/smtp/email",
            headers={"accept": "application/json", "content-type": "application/json",
                     "api-key": settings.brevo_api_key},
            json=_payload, timeout=12,
        )
        if resp.status_code in (200, 201):
            return True
        logger.warning("[EMAIL] Brevo status=%d body=%s", resp.status_code, resp.text[:200])
        return False
    except Exception as e:
        logger.warning("[EMAIL] Envoi échoué : %s", e)
        return False


def _generate_2fa_code() -> str:
    import random
    return f"{random.randint(0, 999999):06d}"


def _send_2fa_email(email: str, code: str) -> bool:
    """
    Envoie le code 2FA par email via l'API HTTP Brevo.
    Utilise requests (pas SMTP) — contourne les restrictions de port Railway.
    """
    try:
        import requests as _req
        resp = _req.post(
            "https://api.brevo.com/v3/smtp/email",
            headers={
                "accept":       "application/json",
                "content-type": "application/json",
                "api-key":      settings.brevo_api_key,
            },
            json={
                "sender":      {"name": settings.brevo_sender_name,
                                "email": settings.brevo_sender_email},
                "to":          [{"email": email}],
                "subject":     f"[Facilim] Votre code de connexion : {code}",
                "textContent": (
                    f"Bonjour,\n\n"
                    f"Votre code de connexion Facilim est :\n\n"
                    f"    {code}\n\n"
                    f"Ce code est valable 10 minutes.\n"
                    f"Si vous n'êtes pas à l'origine de cette demande, ignorez ce message.\n\n"
                    f"— L'équipe Facilim"
                ),
            },
            timeout=10,
        )
        if resp.status_code in (200, 201):
            logger.info("[2FA] Code envoyé à %s", email[:4] + "***")
            return True
        logger.warning("[2FA] API Brevo status=%d body=%s", resp.status_code, resp.text[:200])
        return False
    except Exception as e:
        logger.warning("[2FA] Envoi email échoué : %s", e)
        return False


# ── Dédoublonneur de messages WhatsApp ───────────────────────────────────────
# WhatsApp peut renvoyer le même message_id plusieurs fois (retry sur timeout).
# Ce cache en mémoire supprime les doublons pendant 5 minutes.
_DEDUP_LOCK:       threading.Lock      = threading.Lock()
_DEDUP_CACHE:      dict[str, float]    = {}   # {message_id: timestamp_epoch}
_DEDUP_TTL_SEC:    int                 = 300  # 5 min = TTL cache WhatsApp
_DEDUP_MAX_SIZE:   int                 = 2000 # sécurité anti-fuite mémoire


def _is_duplicate_message(message_id: str) -> bool:
    """
    Retourne True si ce message_id a déjà été traité récemment (doublon WhatsApp).
    Thread-safe. Purge automatique des entrées expirées.
    """
    now = time.monotonic()
    with _DEDUP_LOCK:
        # Purge des entrées expirées
        expired = [mid for mid, ts in _DEDUP_CACHE.items() if now - ts > _DEDUP_TTL_SEC]
        for mid in expired:
            del _DEDUP_CACHE[mid]

        # Cap taille maximale (fail-safe)
        if len(_DEDUP_CACHE) >= _DEDUP_MAX_SIZE:
            oldest = sorted(_DEDUP_CACHE, key=_DEDUP_CACHE.__getitem__)[:200]
            for mid in oldest:
                del _DEDUP_CACHE[mid]

        if message_id in _DEDUP_CACHE:
            return True

        _DEDUP_CACHE[message_id] = now
        return False


logging.basicConfig(
    level=getattr(logging, settings.log_level.upper(), logging.INFO),
    format="%(asctime)s | %(levelname)-8s | %(name)s | %(message)s",
)
logger = logging.getLogger("facilim.app")

# ── Initialisation ────────────────────────────────────────────────────────────
# ── Initialisation du chiffrement ────────────────────────────────────────────
# Priorité 1 : ENCRYPTION_KEY (Fernet 32 bytes base64)
# Priorité 2 : dérivé du JWT_SECRET_KEY via SHA-256 (fallback automatique)
# Priorité 3 : pas de chiffrement (warning — ne bloque pas le démarrage)
import os as _os, hashlib as _hashlib, base64 as _base64

_enc_key = settings.encryption_key.strip()
if not _enc_key:
    _jwt = _os.environ.get("JWT_SECRET_KEY", "").strip()
    if _jwt:
        _derived    = _hashlib.sha256(_jwt.encode()).digest()
        _enc_key    = _base64.urlsafe_b64encode(_derived).decode()
        logger.info("[STARTUP] ENCRYPTION_KEY absent — clé dérivée depuis JWT_SECRET_KEY.")

try:
    encryption.initialize(_enc_key)
except Exception as _e:
    logger.warning("[STARTUP] Chiffrement non actif : %s", _e)
db_conn_module.configure(settings.database_url)
db_conn_module.initialize_schema()

# ── Migrations idempotentes — s'exécutent au démarrage ───────────────────────
try:
    from app.database.migrations import run_all as _run_migrations
    _migration_result = _run_migrations(verbose=False)
    if _migration_result.get("applied"):
        logger.info("[STARTUP] Migrations appliquées : %s", _migration_result["applied"])
    if _migration_result.get("errors"):
        logger.error("[STARTUP] Erreurs migrations : %s", _migration_result["errors"])
except Exception as _mig_err:
    logger.warning("[STARTUP] Migrations non bloquantes : %s", _mig_err)

# ── Vérification et correction des colonnes Phase 3+ ─────────────────────────
# SQLite peut marquer une migration comme appliquée même si certaines ALTER TABLE
# ont échoué silencieusement (contention, verrou). Ce bloc garantit que toutes
# les colonnes nécessaires existent, quelle que soit l'historique des migrations.
_COLONNES_PHASE3 = [
    ("profil_principal",          "TEXT DEFAULT ''"),
    ("profil_secondaire",         "TEXT DEFAULT ''"),
    ("tags_detectes",             "TEXT DEFAULT '[]'"),
    ("texte_b_vie_quotidienne",   "TEXT DEFAULT ''"),
    ("texte_c_scolarite",         "TEXT DEFAULT ''"),
    ("texte_d_situation_pro",     "TEXT DEFAULT ''"),
    ("texte_e_projet_vie",        "TEXT DEFAULT ''"),
    ("score_maturite_cerfa",      "INTEGER DEFAULT 0"),
    ("niveau_maturite",           "TEXT DEFAULT ''"),
    ("alertes_qualite_json",      "TEXT DEFAULT '[]'"),
    ("rapport_qualite_json",      "TEXT DEFAULT '{}'"),
    ("analyse_situation_json",    "TEXT DEFAULT '{}'"),
    ("synthese_situation",        "TEXT DEFAULT ''"),
    ("niveau_confiance_analyse",  "TEXT DEFAULT ''"),
    ("cockpit_pro_json",          "TEXT DEFAULT '{}'"),  # FACILIM PROD-1 : cockpit consolidé
    ("version",                   "INTEGER DEFAULT 1"),  # FIX-TEST4 : ALTER migration 008 avalé par le découpeur
]
try:
    _startup_db = db_conn_module._get_raw_connection()
    _cols_existantes = {r[1] for r in _startup_db.execute("PRAGMA table_info(dossiers)").fetchall()}
    _ajoutees = []
    for _col, _typedef in _COLONNES_PHASE3:
        if _col not in _cols_existantes:
            try:
                _startup_db.execute(f"ALTER TABLE dossiers ADD COLUMN {_col} {_typedef}")
                _startup_db.commit()
                _ajoutees.append(_col)
            except Exception as _ce:
                logger.warning("[STARTUP] Colonne %s non ajoutée : %s", _col, _ce)
    if _ajoutees:
        logger.info("[STARTUP] Colonnes Phase 3+ ajoutées : %s", _ajoutees)
    else:
        logger.info("[STARTUP] Colonnes Phase 3+ : toutes présentes ✓")
    _startup_db.close()
except Exception as _col_err:
    logger.warning("[STARTUP] Vérification colonnes non bloquante : %s", _col_err)

# ── Client LLM ───────────────────────────────────────────────────────────────
from openai import OpenAI
_llm_client = OpenAI(api_key=settings.openai_api_key)

# ── Observabilité Sentry — suivi des erreurs en production ───────────────────
# Activé uniquement si un DSN est fourni. Garde-fou RGPD : on retire les corps de
# requête, les cookies et les en-têtes sensibles pour ne JAMAIS transmettre de
# donnée de santé ou personnelle brute à un tiers.
if settings.sentry_dsn:
    try:
        import sentry_sdk

        def _sentry_scrub(event, hint):
            req = event.get("request") or {}
            req.pop("data", None)
            req.pop("cookies", None)
            req.pop("query_string", None)
            if isinstance(req.get("headers"), dict):
                req["headers"] = {
                    k: v for k, v in req["headers"].items()
                    if k.lower() not in ("authorization", "cookie", "api-key")
                }
            if req:
                event["request"] = req
            return event

        sentry_sdk.init(
            dsn=settings.sentry_dsn,
            environment=settings.sentry_environment,
            send_default_pii=False,      # jamais d'IP, de cookies ou de corps par défaut
            traces_sample_rate=0.0,      # pas de tracing perf, uniquement les erreurs
            before_send=_sentry_scrub,
        )
        logger.info("[SENTRY] Suivi des erreurs activé (env=%s)", settings.sentry_environment)
    except Exception as _sentry_err:
        logger.warning("[SENTRY] Initialisation ignorée (non bloquant) : %s", _sentry_err)

# ── Services ─────────────────────────────────────────────────────────────────
_wa_service = WhatsAppService(
    api_token=settings.whatsapp_api_token,
    phone_number_id=settings.whatsapp_phone_number_id,
    api_version=settings.whatsapp_api_version,
)
_notif_service = NotificationService(whatsapp_service=_wa_service)

# ── FastAPI ───────────────────────────────────────────────────────────────────
app = FastAPI(
    title="Facilim API v2",
    description=(
        "Infrastructure nationale de sécurisation et d'orchestration "
        "des parcours MDPH en France."
    ),
    version="2.0.0",
    docs_url="/docs" if not settings.is_production else None,
    redoc_url=None,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["https://facilim.fr"] if settings.is_production else ["*"],
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "PATCH", "DELETE"],
    allow_headers=["*"],
)

# Fichiers statiques (dashboard, login)
import os
_static_dir = os.path.join(os.path.dirname(__file__), "..", "static")
if os.path.isdir(_static_dir):
    app.mount("/static", StaticFiles(directory=_static_dir), name="static")

_dashboards_dir = os.path.join(os.path.dirname(__file__), "dashboards")
if os.path.isdir(_dashboards_dir):
    app.mount("/dashboards", StaticFiles(directory=_dashboards_dir), name="dashboards")


# ── Dépendances ───────────────────────────────────────────────────────────────

def _get_db():
    """Ouvre une connexion DB par requête (context manager non utilisable en Depends)."""
    conn = db_conn_module._get_raw_connection()
    try:
        yield conn
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


def _get_current_user(request: Request) -> dict[str, Any]:
    """Vérifie le JWT du Dashboard. Lève 401 si invalide."""
    token = request.cookies.get("facilim_token") or (
        request.headers.get("Authorization", "").removeprefix("Bearer ").strip() or None
    )
    if not token:
        raise HTTPException(status_code=401, detail="Non authentifié")
    payload = token_manager.verify_dashboard_token(token, settings.jwt_secret_key)
    if not payload:
        raise HTTPException(status_code=401, detail="Token invalide ou expiré")
    return payload


def _get_orchestrator(db=Depends(_get_db)) -> OrchestrationEngine:
    return OrchestrationEngine(
        db_conn=db,
        whatsapp_service=_wa_service,
        notification_service=_notif_service,
        openai_client=_llm_client,
        settings=settings,
    )


# ── Health ────────────────────────────────────────────────────────────────────

@app.get("/health")
def health(db=Depends(_get_db)):
    """
    Endpoint de santé — utilisé par Railway et les outils de monitoring.
    Vérifie les variables d'environnement et la base de données.
    OpenAI non testé ici (évite le coût à chaque sonde).
    """
    from app.engines.health_check import health_check
    report = health_check(db_conn=db, check_openai=False)
    status_code = 200 if report.status in ("OK", "DEGRADED") else 503
    return JSONResponse(
        status_code=status_code,
        content={
            "status":    report.status,
            "version":   "2.0.0",
            "env":       settings.app_env,
            "db":        report.db_ok,
            "warnings":  report.warnings,
            "errors":    report.errors,
            "timestamp": report.timestamp,
        },
    )


@app.get("/api/v1/health/full")
def health_full(user=Depends(_get_current_user), db=Depends(_get_db)):
    """
    Rapport de santé complet avec test OpenAI — réservé aux administrateurs.
    """
    from app.engines.health_check import health_check
    report = health_check(db_conn=db, check_openai=True)
    return report.to_dict()


# ── Webhook WhatsApp ──────────────────────────────────────────────────────────

@app.get("/api/v1/webhook/whatsapp")
async def whatsapp_verify(request: Request):
    """Vérification du webhook WhatsApp (challenge Meta)."""
    params = request.query_params
    if (
        params.get("hub.mode") == "subscribe"
        and params.get("hub.verify_token") == settings.whatsapp_verify_token
    ):
        return Response(content=params.get("hub.challenge", ""), media_type="text/plain")
    raise HTTPException(status_code=403, detail="Token de vérification invalide")


@app.post("/api/v1/webhook/whatsapp")
async def whatsapp_webhook(
    request: Request,
    background_tasks: BackgroundTasks,
):
    """
    Réception des messages WhatsApp entrants.

    Acquitte immédiatement un 200 OK à WhatsApp (avant le timeout de 5 s),
    puis délègue le traitement complet en tâche de fond.
    """
    try:
        payload = await request.json()
    except Exception:
        return JSONResponse({"status": "ok"})

    messages = parse_webhook_payload(payload)
    for msg in messages:
        if msg.get("type") in ("text", "image", "document", "audio"):
            mid = msg["message_id"]
            if _is_duplicate_message(mid):
                logger.info(f"[WA] Doublon ignoré : {mid[:12]}…")
                continue
            _wa_service.mark_as_read(mid)
            background_tasks.add_task(
                _process_whatsapp_background,
                msg["from"],
                msg.get("text", ""),
                mid,
                msg.get("media_id"),
                msg.get("mime_type"),
            )

    return JSONResponse({"status": "ok"})


def _process_whatsapp_background(
    from_number: str,
    message_text: str,
    message_id: str,
    media_id: str | None,
    mime_type: str | None,
) -> None:
    """
    Exécuté en arrière-plan par FastAPI (thread pool).
    Ouvre sa propre connexion DB — indépendante du cycle de vie de la requête HTTP.
    """
    conn = db_conn_module._get_raw_connection()
    try:
        orchestrator = OrchestrationEngine(
            db_conn=conn,
            whatsapp_service=_wa_service,
            notification_service=_notif_service,
            openai_client=_llm_client,
            settings=settings,
        )
        orchestrator.handle_whatsapp_message(
            from_number=from_number,
            message_text=message_text,
            message_id=message_id,
            media_id=media_id,
            mime_type=mime_type,
        )
        conn.commit()
    except Exception as e:
        logger.error(f"[BG/WA] Erreur traitement message {message_id[:8]} : {e}", exc_info=True)
        conn.rollback()
    finally:
        conn.close()


# ── Formulaire de contact (demandes de renseignement des structures) ──────────
class ContactRequest(BaseModel):
    prenom:    str
    nom:       str
    email:     str
    structure: str | None = ""
    type:      str | None = ""
    message:   str | None = ""


@app.post("/api/v1/contact")
def contact_form(body: ContactRequest):
    """
    Reçoit une demande de renseignement depuis la page d'accueil et l'achemine
    par email vers la boîte de contact Facilim. La réponse part directement vers
    le prospect (replyTo). Aucune donnée de santé ici : prospection uniquement.
    """
    import html as _html
    _esc = lambda s: _html.escape(str(s or "").strip())
    prenom, nom, email = _esc(body.prenom), _esc(body.nom), _esc(body.email)
    if not (prenom and nom and email and "@" in email):
        raise HTTPException(status_code=400, detail="Prénom, nom et email valides requis.")

    corps = (
        f"<p><strong>Nouvelle demande via facilim.fr</strong></p>"
        f"<p><b>Nom :</b> {prenom} {nom}<br/>"
        f"<b>Email :</b> {email}<br/>"
        f"<b>Structure :</b> {_esc(body.structure) or '—'}<br/>"
        f"<b>Type :</b> {_esc(body.type) or '—'}</p>"
        f"<p><b>Message :</b><br/>{_esc(body.message).replace(chr(10), '<br/>') or '—'}</p>"
    )
    envoye = _send_email_brevo(
        to_email=settings.contact_email,
        subject=f"[Facilim] Demande de {prenom} {nom}" + (f" — {_esc(body.structure)}" if body.structure else ""),
        html=corps,
        reply_to_email=email,
        reply_to_name=f"{prenom} {nom}",
    )
    # Accusé de réception au prospect (non bloquant).
    if envoye:
        _send_email_brevo(
            to_email=email,
            subject="Votre demande a bien été reçue — Facilim",
            html=(f"<p>Bonjour {prenom},</p><p>Merci pour votre message. "
                  f"Nous revenons vers vous très vite.</p><p>L'équipe Facilim</p>"),
        )
    logger.info("[CONTACT] demande reçue de %s (envoi=%s)", email[:4] + "***", envoye)
    return {"success": True, "envoye": envoye}


# ── Documents reçus par email → rattachement automatique au dossier ───────────
def _ingerer_document_dans_dossier(db, dossier_id: str, content: bytes,
                                   filename: str, content_type: str) -> bool:
    """
    Ingestion d'un document reçu (email, upload) dans un dossier : extraction du
    texte, extraction LLM des champs utiles (sans rien écraser), enregistrement
    comme pièce justificative. Réutilise la même logique que l'upload tableau de bord.
    """
    texte = _extraire_texte_fichier(content, filename, content_type)
    extraits: dict = {}
    if texte and len(texte) > 50 and not texte.startswith("["):
        try:
            prompt = f"""Extrais les informations MDPH de ce document. Retourne UNIQUEMENT un JSON
avec les champs trouvés parmi : nom_prenom, date_naissance, genre, commune_naissance,
pays_naissance, adresse_complete, departement, num_secu, diagnostics, traitements,
medecin_traitant, impact_quotidien, statut_emploi, historique_mdph, restrictions_emploi.
Règles : genre déduit de la civilité (Mr/M./Monsieur → "homme" ; Mme/Madame → "femme").
N'invente rien ; si un champ est absent, ne l'inclus pas.

Document :
{texte[:4000]}"""
            resp = _llm_client.chat.completions.create(
                model=settings.openai_model_fast,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=400, temperature=0.0,
                response_format={"type": "json_object"},
            )
            extraits = json.loads(resp.choices[0].message.content) or {}
        except Exception as e:
            logger.warning("[EMAIL_IN] extraction LLM ignorée : %s", e)

    row = db.execute("SELECT synthese_json FROM dossiers WHERE id = ?", (dossier_id,)).fetchone()
    synthese = json.loads((row["synthese_json"] if row else "{}") or "{}")
    for k, v in extraits.items():
        # Ne pas écraser une valeur existante ; ignorer les gabarits de date.
        if not v or synthese.get(k):
            continue
        if k == "date_naissance" and not any(c.isdigit() for c in str(v)):
            continue
        synthese[k] = v
    if texte and not texte.startswith("["):
        _notes = synthese.get("documents_texte", "")
        synthese["documents_texte"] = f"{_notes}\n\n--- {filename} (reçu par email) ---\n{texte[:2000]}"

    now = datetime.now(timezone.utc).isoformat()
    db.execute("UPDATE dossiers SET synthese_json = ?, updated_at = ? WHERE id = ?",
               (json.dumps(synthese, ensure_ascii=False), now, dossier_id))
    db.execute(
        """INSERT INTO pieces_justificatives
           (id, dossier_id, type_piece, mime_type, uploaded_par, ocr_effectue,
            flag_validation_humaine, created_at, updated_at)
           VALUES (?, ?, 'DOCUMENT_EMAIL', ?, 'email', 0, 0, ?, ?)""",
        (str(uuid.uuid4()), dossier_id, content_type or "", now, now),
    )
    return True


@app.post("/api/v1/webhook/email-inbound")
async def email_inbound_webhook(request: Request, db=Depends(_get_db)):
    """
    Webhook d'email entrant (ex. Brevo Inbound Parsing). Quand un usager envoie des
    documents par email à l'adresse Facilim, ses pièces jointes sont rattachées
    automatiquement à SON dossier — le plus récent associé à son adresse.

    Sécurité : jeton partagé requis (?token= ou en-tête X-Inbound-Token = INBOUND_EMAIL_SECRET).
    """
    token = request.query_params.get("token") or request.headers.get("x-inbound-token", "")
    if not settings.inbound_email_secret or token != settings.inbound_email_secret:
        raise HTTPException(status_code=403, detail="Webhook non autorisé.")
    try:
        payload = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Payload JSON invalide.")

    items = payload.get("items") if isinstance(payload, dict) else None
    items = items if isinstance(items, list) else [payload]
    import base64 as _b64
    traites = 0
    for item in items:
        if not isinstance(item, dict):
            continue
        sender = ((item.get("From") or {}).get("Address") if isinstance(item.get("From"), dict)
                  else item.get("from") or item.get("sender") or "")
        sender = str(sender or "").strip().lower()
        attachments = item.get("Attachments") or item.get("attachments") or []
        if not sender or not attachments:
            continue
        # Rattachement par email, dossier le plus récent (anti-fuite : jamais le 1er au hasard).
        # Anti-fuite renforcé : l'adresse doit apparaître comme VALEUR JSON exacte
        # ("email@x.fr"), jamais comme simple sous-chaîne d'un texte libre — sinon une
        # note pro citant un email rattacherait des pièces médicales au mauvais dossier.
        # On échappe aussi les métacaractères LIKE ('_' et '%') du local-part, qui
        # élargiraient silencieusement la correspondance. Un email non rattaché est
        # simplement loggé/ignoré : mode d'échec bien plus sûr qu'une fuite.
        _sender_like = sender.replace("\\", "\\\\").replace("%", "\\%").replace("_", "\\_")
        row = db.execute(
            "SELECT id FROM dossiers WHERE lower(synthese_json) LIKE ? ESCAPE '\\' "
            "ORDER BY updated_at DESC LIMIT 1",
            (f'%"{_sender_like}"%',),
        ).fetchone()
        if not row:
            logger.info("[EMAIL_IN] aucun dossier pour l'expéditeur — pièces ignorées")
            continue
        dossier_id = row["id"]
        for att in (attachments if isinstance(attachments, list) else []):
            try:
                name  = att.get("Name") or att.get("name") or "document"
                ctype = att.get("ContentType") or att.get("contentType") or ""
                c_b64 = att.get("Content") or att.get("content") or ""
                if not c_b64:
                    continue
                _ingerer_document_dans_dossier(db, dossier_id, _b64.b64decode(c_b64), name, ctype)
                traites += 1
            except Exception as _att_err:
                logger.warning("[EMAIL_IN] pièce ignorée : %s", _att_err)
        try:
            db.commit()
        except Exception:
            pass
    logger.info("[EMAIL_IN] %d pièce(s) rattachée(s)", traites)
    return {"success": True, "pieces_traitees": traites}


# ── Authentification Dashboard ────────────────────────────────────────────────

# ── Rate limiting authentification ───────────────────────────────────────────
# Protection brute-force : 5 tentatives max par adresse email sur 5 minutes.
_login_attempts:   dict[str, list[float]] = {}
_login_lock        = threading.Lock()
_LOGIN_MAX         = 5
_LOGIN_WINDOW_SEC  = 300   # 5 minutes


def _check_login_rate(email: str) -> bool:
    """Retourne False si l'email est temporairement bloqué."""
    now = time.time()
    with _login_lock:
        attempts = [t for t in _login_attempts.get(email, []) if now - t < _LOGIN_WINDOW_SEC]
        if len(attempts) >= _LOGIN_MAX:
            return False
        attempts.append(now)
        _login_attempts[email] = attempts
        return True


class LoginRequest(BaseModel):
    email: str
    password: str


@app.post("/auth/login")         # alias legacy (ancien frontend)
@app.post("/api/v1/auth/login")
def login(body: LoginRequest, response: Response, db=Depends(_get_db)):
    """Authentification Dashboard ESSMS."""
    # Rate limiting — 5 tentatives / 5 min par email
    if not _check_login_rate(body.email):
        raise HTTPException(
            status_code=429,
            detail="Trop de tentatives. Réessayez dans 5 minutes.",
        )

    # Vérification des credentials
    # Priorité 1 : AUTH_PASSWORD_HASH (bcrypt) — sécurisé
    # Priorité 2 : AUTH_PASSWORD (plaintext) — rétrocompatibilité Railway
    import bcrypt as _bcrypt
    import os as _os

    def _verify_account(email_cfg: str, pw_hash_cfg: str, pw_plain_cfg: str) -> bool:
        """Vérifie un compte : bcrypt si hash présent, sinon plaintext (legacy)."""
        if not email_cfg or body.email != email_cfg:
            return False
        if pw_hash_cfg:
            try:
                return _bcrypt.checkpw(body.password.encode(), pw_hash_cfg.encode())
            except Exception:
                pass
        if pw_plain_cfg:
            # Fallback plaintext (Railway legacy — à supprimer après migration)
            if pw_plain_cfg == body.password:
                logger.warning(
                    "[AUTH] Connexion via mot de passe en clair pour %s — "
                    "configurer AUTH_PASSWORD_HASH dans Railway.",
                    email_cfg,
                )
                return True
        return False

    # Lire les anciens mots de passe plaintext depuis l'environnement (fallback)
    _old_pw1 = _os.environ.get("AUTH_PASSWORD", "")
    _old_pw2 = _os.environ.get("AUTH_PASSWORD_2", "")

    matched = (
        _verify_account(settings.auth_email,   settings.auth_password_hash,  _old_pw1)
        or
        _verify_account(settings.auth_email_2, settings.auth_password_hash_2, _old_pw2)
    )
    if not matched:
        raise HTTPException(status_code=401, detail="Identifiants incorrects")

    # Récupération ou création de l'utilisateur
    row = db.execute(
        "SELECT * FROM utilisateurs WHERE email = ? AND actif = 1 LIMIT 1",
        (body.email,),
    ).fetchone()

    if row:
        user = dict(row)
    else:
        user_id = str(uuid.uuid4())
        now = datetime.now(timezone.utc).isoformat()
        db.execute(
            """
            INSERT OR IGNORE INTO utilisateurs
                (id, email, password_hash, role, actif, created_at, updated_at)
            VALUES (?, ?, '[LEGACY]', 'EDUCATEUR', 1, ?, ?)
            """,
            (user_id, body.email, now, now),
        )
        user = {"id": user_id, "email": body.email, "role": "EDUCATEUR", "organisation_id": None}

    # Mise à jour dernière connexion
    db.execute(
        "UPDATE utilisateurs SET derniere_connexion = ? WHERE email = ?",
        (datetime.now(timezone.utc).isoformat(), body.email),
    )

    token, expire_at = token_manager.create_dashboard_token(
        utilisateur_id=user["id"],
        email=user["email"],
        role=user.get("role", "EDUCATEUR"),
        organisation_id=user.get("organisation_id"),
        secret_key=settings.jwt_secret_key,
        expire_minutes=settings.jwt_expire_minutes,
    )

    event_logger.log_event(
        "CONNEXION_DASHBOARD",
        utilisateur_id=user["id"],
        canal="dashboard",
        db_conn=db,
    )

    # ── 2FA : génère et envoie le code, stocke le JWT en attente ─────────────
    code = _generate_2fa_code()
    expires = time.time() + _2FA_TTL_SEC
    with _2FA_LOCK:
        _2FA_STORE[body.email] = {"code": code, "expires": expires, "token": token,
                                  "user": user, "expire_minutes": settings.jwt_expire_minutes}

    sent = _send_2fa_email(body.email, code)
    if not sent:
        # Fallback : si l'email échoue, log le code (dev/staging uniquement)
        logger.warning("[2FA] Email non envoyé — code=%s pour %s", code, body.email[:4] + "***")

    # Ne pas poser le cookie ici — il sera posé après vérification du code
    return {"success": True, "email": user["email"]}


@app.get("/verify")
def verify_page():
    """Page de saisie du code 2FA."""
    verify_path = os.path.join(os.path.dirname(__file__), "..", "static", "verify.html")
    if os.path.isfile(verify_path):
        return FileResponse(verify_path, headers={"Cache-Control": "no-cache, no-store, must-revalidate"})
    return HTMLResponse("<h1>Page de vérification non trouvée</h1>", status_code=404)


class VerifyRequest(BaseModel):
    email: str
    code:  str


@app.post("/auth/verify")
@app.post("/api/v1/auth/verify")
def verify_code(body: VerifyRequest, response: Response):
    """Validation du code 2FA — pose le cookie JWT si correct."""
    now = time.time()
    with _2FA_LOCK:
        entry = _2FA_STORE.get(body.email)

    if not entry:
        raise HTTPException(status_code=400, detail="Aucun code en attente pour cet email.")
    if now > entry["expires"]:
        with _2FA_LOCK:
            _2FA_STORE.pop(body.email, None)
        raise HTTPException(status_code=400, detail="Code expiré. Reconnectez-vous.")
    if entry["code"] != body.code.strip():
        raise HTTPException(status_code=401, detail="Code incorrect — réessayez.")

    # Code correct → poser le cookie JWT et purger le store
    with _2FA_LOCK:
        _2FA_STORE.pop(body.email, None)

    user           = entry["user"]
    token          = entry["token"]
    expire_minutes = entry["expire_minutes"]

    response.set_cookie(
        key="facilim_token",
        value=token,
        httponly=True,
        secure=settings.is_production,
        samesite="lax",
        max_age=expire_minutes * 60,
    )
    return {"success": True, "role": user.get("role", "EDUCATEUR"), "email": user["email"]}


@app.post("/api/v1/auth/logout")
def logout(response: Response):
    response.delete_cookie("facilim_token")
    return {"success": True}


@app.get("/logout")
def logout_redirect(response: Response):
    response.delete_cookie("facilim_token")
    return RedirectResponse("/login", status_code=302)


# ── Dashboard — Dossiers ──────────────────────────────────────────────────────

@app.get("/api/v1/dashboard/dossiers")
def list_dossiers(
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """Liste des dossiers visibles par l'utilisateur connecté."""
    event_logger.log_event("ACCES_DASHBOARD", utilisateur_id=user.get("sub"), canal="dashboard", db_conn=db)

    if user.get("role") == "SUPER_ADMIN":
        rows = db.execute(
            "SELECT * FROM dossiers WHERE deleted_at IS NULL ORDER BY updated_at DESC LIMIT 200"
        ).fetchall()
    else:
        org_id = user.get("org_id")
        rows = db.execute(
            "SELECT * FROM dossiers WHERE organisation_id = ? AND deleted_at IS NULL "
            "ORDER BY updated_at DESC LIMIT 200",
            (org_id,),
        ).fetchall() if org_id else db.execute(
            "SELECT * FROM dossiers WHERE deleted_at IS NULL ORDER BY updated_at DESC LIMIT 200"
        ).fetchall()

    dossiers = []
    for row in rows:
        d = dict(row)
        d["droits_identifies"] = json.loads(d.get("droits_identifies_json", "[]") or "[]")
        d.pop("droits_identifies_json", None)
        dossiers.append(d)

    return {"dossiers": dossiers, "total": len(dossiers)}


@app.get("/api/v1/dashboard/dossiers/{dossier_id}")
def get_dossier(
    dossier_id: str,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """Détail d'un dossier avec consentement et scoring."""
    # deleted_at IS NULL : un dossier supprimé ne doit plus pouvoir être rouvert
    # (sinon une alerte résiduelle le fait « revenir »).
    row = db.execute(
        "SELECT * FROM dossiers WHERE id = ? AND deleted_at IS NULL", (dossier_id,)
    ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Dossier non trouvé ou supprimé")

    d = dict(row)
    access_control.require_permission(user.get("role", "LECTEUR"), access_control.Permission.DOSSIER_READ)

    d["droits_identifies"] = json.loads(d.get("droits_identifies_json", "[]") or "[]")
    d["synthese"] = json.loads(d.get("synthese_json", "{}") or "{}")
    d["questions"] = json.loads(d.get("questions_json", "[]") or "[]")
    d["conversation"] = json.loads(d.get("conversation_json", "[]") or "[]")
    # Rapport qualité (inclut les signaux experts de la revue instructeur :
    # robustesse, risque de refus, axes d'amélioration) pour le cockpit.
    d["rapport_qualite"] = json.loads(d.get("rapport_qualite_json", "{}") or "{}")
    d["revue_instructeur"] = d["rapport_qualite"].get("revue_instructeur", {})

    # CÂBLAGE P3/P4 — complétude DIFFÉRENCIÉE (5 états par champ) + verdict métier
    # « prêt à transmettre ». Réutilise synthese_completude + validite_dossier déjà
    # développés : le professionnel voit ce qui manque, ce qu'il doit compléter, ce
    # qui a été refusé, et POURQUOI le dossier n'est pas prêt. Ne recrée aucun score.
    try:
        from app.services.conversation.router import get_agent
        from app.engines.revue_instructeur import validite_dossier
        _profil = (d["synthese"].get("profil_mdph") or d.get("service_type")
                   or d["synthese"].get("service_type") or "adulte")
        _ag = get_agent(_profil)
        d["completude"] = _ag.synthese_completude(d["synthese"])
        d["completude"]["etats_champs"] = _ag.etat_par_champ(d["synthese"])
        d["validite_dossier"] = (d["rapport_qualite"].get("validite_dossier")
                                 or validite_dossier(d["synthese"], _profil))
    except Exception as _compl_err:
        logger.warning("[COMPLETUDE] synthèse non calculée pour %s : %s", dossier_id, _compl_err)

    # Consentement
    usager_row = db.execute("SELECT id FROM usagers WHERE id = ?", (d.get("usager_id", ""),)).fetchone()
    if usager_row:
        from app.services.consent_service import get_consent_summary
        d["consent"] = get_consent_summary(usager_row["id"], db)

    # Pièces
    pieces = db.execute(
        "SELECT id, type_piece, ocr_effectue, score_confiance_ocr, flag_validation_humaine, "
        "validee_par, created_at FROM pieces_justificatives WHERE dossier_id = ?",
        (dossier_id,),
    ).fetchall()
    d["pieces"] = [dict(p) for p in pieces]

    # Scoring
    scoring = db.execute(
        "SELECT * FROM analyses_scoring WHERE dossier_id = ? ORDER BY created_at DESC LIMIT 1",
        (dossier_id,),
    ).fetchone()
    if scoring:
        d["scoring"] = dict(scoring)

    event_logger.log_event(
        "ACCES_DOSSIER",
        dossier_id=dossier_id,
        utilisateur_id=user.get("sub"),
        canal="dashboard",
        db_conn=db,
    )
    return d


@app.get("/api/v1/dashboard/metriques")
def get_metriques(user=Depends(_get_current_user), db=Depends(_get_db)):
    """
    Métriques de qualité et d'utilisation pour le pilotage de Facilim.

    Trois catégories :
      - qualite   : enrichissement, alertes, maturité, confiance
      - utilisation : dossiers, durée, relances, abandons
      - analyse   : risques, protections, acteurs
    """
    # ── Qualité ───────────────────────────────────────────────────────────────
    # Requête résiliente : vérifie d'abord les colonnes disponibles pour éviter
    # une erreur 500 si les colonnes Phase 3+ ne sont pas encore présentes.
    try:
        _cols_dossiers = {r[1] for r in db.execute("PRAGMA table_info(dossiers)").fetchall()}
        _has_maturite  = "score_maturite_cerfa" in _cols_dossiers
        _has_confiance = "niveau_confiance_analyse" in _cols_dossiers
        _has_niveau    = "niveau_maturite" in _cols_dossiers

        _select_maturite  = "AVG(score_maturite_cerfa)" if _has_maturite  else "0"
        _select_nb_faible = "COUNT(CASE WHEN niveau_maturite = 'FAIBLE'    THEN 1 END)" if _has_niveau else "0"
        _select_nb_moyen  = "COUNT(CASE WHEN niveau_maturite = 'MOYEN'     THEN 1 END)" if _has_niveau else "0"
        _select_nb_solide = "COUNT(CASE WHEN niveau_maturite = 'SOLIDE'    THEN 1 END)" if _has_niveau else "0"
        _select_nb_excel  = "COUNT(CASE WHEN niveau_maturite = 'EXCELLENT' THEN 1 END)" if _has_niveau else "0"
        _select_conf_fort = "COUNT(CASE WHEN niveau_confiance_analyse = 'FORT'   THEN 1 END)" if _has_confiance else "0"
        _select_conf_moy  = "COUNT(CASE WHEN niveau_confiance_analyse = 'MOYEN'  THEN 1 END)" if _has_confiance else "0"
        _select_conf_fai  = "COUNT(CASE WHEN niveau_confiance_analyse = 'FAIBLE' THEN 1 END)" if _has_confiance else "0"

        stats_qualite = db.execute(f"""
            SELECT
                COUNT(*) as nb_total,
                {_select_maturite}  as maturite_moy,
                {_select_nb_faible} as nb_faible,
                {_select_nb_moyen}  as nb_moyen,
                {_select_nb_solide} as nb_solide,
                {_select_nb_excel}  as nb_excellent,
                {_select_conf_fort} as nb_confiance_fort,
                {_select_conf_moy}  as nb_confiance_moyen,
                {_select_conf_fai}  as nb_confiance_faible
            FROM dossiers WHERE deleted_at IS NULL
        """).fetchone()
    except Exception:
        stats_qualite = None

    # Taux d'enrichissement moyen (dans synthese_json)
    enrichissements = db.execute("""
        SELECT synthese_json FROM dossiers
        WHERE synthese_json LIKE '%_taux_enrichissement_moyen%' AND deleted_at IS NULL
        LIMIT 200
    """).fetchall()
    taux_enrichissement_list = []
    for row in enrichissements:
        try:
            d = json.loads(row["synthese_json"] or "{}")
            t = d.get("_taux_enrichissement_moyen")
            if t and isinstance(t, (int, float)):
                taux_enrichissement_list.append(t)
        except Exception:
            pass
    taux_enrichissement_moy = (
        round(sum(taux_enrichissement_list) / len(taux_enrichissement_list), 2)
        if taux_enrichissement_list else None
    )

    # Alertes qualité récurrentes (top 5 messages)
    alertes_rows = db.execute("""
        SELECT alertes_qualite_json FROM dossiers
        WHERE alertes_qualite_json != '[]' AND alertes_qualite_json IS NOT NULL
        AND deleted_at IS NULL LIMIT 500
    """).fetchall()
    alertes_compteur: dict[str, int] = {}
    for row in alertes_rows:
        try:
            alertes = json.loads(row["alertes_qualite_json"] or "[]")
            for a in alertes:
                cle = str(a)[:80]
                alertes_compteur[cle] = alertes_compteur.get(cle, 0) + 1
        except Exception:
            pass
    alertes_top = sorted(alertes_compteur.items(), key=lambda x: -x[1])[:5]

    # ── Utilisation ───────────────────────────────────────────────────────────
    stats_util = db.execute("""
        SELECT
            COUNT(*) as nb_total,
            COUNT(CASE WHEN statut = 'EN_COLLECTE'  THEN 1 END) as nb_collecte,
            COUNT(CASE WHEN statut = 'EN_ANALYSE'   THEN 1 END) as nb_analyse,
            COUNT(CASE WHEN statut = 'COMPLET'      THEN 1 END) as nb_complet,
            COUNT(CASE WHEN statut = 'INCOMPLET'    THEN 1 END) as nb_incomplet,
            COUNT(CASE WHEN statut = 'SOUMIS'       THEN 1 END) as nb_soumis
        FROM dossiers WHERE deleted_at IS NULL
    """).fetchone()

    # Relances (comptage dans synthese_json)
    relances_rows = db.execute("""
        SELECT synthese_json FROM dossiers
        WHERE synthese_json LIKE '%_relances_faites%' AND deleted_at IS NULL
        LIMIT 200
    """).fetchall()
    nb_relances_total = 0
    nb_dossiers_avec_relance = 0
    for row in relances_rows:
        try:
            d = json.loads(row["synthese_json"] or "{}")
            rf = d.get("_relances_faites") or {}
            if rf:
                nb_dossiers_avec_relance += 1
                nb_relances_total += len(rf)
        except Exception:
            pass

    # ── Analyse ───────────────────────────────────────────────────────────────
    analyse_rows = db.execute("""
        SELECT analyse_situation_json FROM dossiers
        WHERE analyse_situation_json != '{}' AND analyse_situation_json IS NOT NULL
        AND deleted_at IS NULL LIMIT 200
    """).fetchall()

    risques_compteur: dict[str, int] = {}
    protections_compteur: dict[str, int] = {}
    acteurs_compteur: dict[str, int] = {}

    for row in analyse_rows:
        try:
            a = json.loads(row["analyse_situation_json"] or "{}")
            for r in a.get("facteurs_risque", []):
                risques_compteur[str(r)[:60]] = risques_compteur.get(str(r)[:60], 0) + 1
            for p in a.get("facteurs_protection", []):
                protections_compteur[str(p)[:60]] = protections_compteur.get(str(p)[:60], 0) + 1
            for act in a.get("acteurs_mobilisables", []):
                nom = act.get("acteur", "")
                if nom:
                    acteurs_compteur[nom] = acteurs_compteur.get(nom, 0) + 1
        except Exception:
            pass

    return {
        "qualite": {
            "nb_dossiers_total":       stats_qualite["nb_total"] if stats_qualite else 0,
            "maturite_moyenne":        round(stats_qualite["maturite_moy"] or 0, 1) if stats_qualite else 0,
            "repartition_maturite": {
                "faible":    stats_qualite["nb_faible"]    if stats_qualite else 0,
                "moyen":     stats_qualite["nb_moyen"]     if stats_qualite else 0,
                "solide":    stats_qualite["nb_solide"]    if stats_qualite else 0,
                "excellent": stats_qualite["nb_excellent"] if stats_qualite else 0,
            },
            "repartition_confiance": {
                "fort":   stats_qualite["nb_confiance_fort"]   if stats_qualite else 0,
                "moyen":  stats_qualite["nb_confiance_moyen"]  if stats_qualite else 0,
                "faible": stats_qualite["nb_confiance_faible"] if stats_qualite else 0,
            },
            "taux_enrichissement_moyen": taux_enrichissement_moy,
            "alertes_recurrentes":       [{"message": k, "occurrences": v} for k, v in alertes_top],
        },
        "utilisation": {
            "nb_dossiers_total":        stats_util["nb_total"]     if stats_util else 0,
            "en_collecte":              stats_util["nb_collecte"]  if stats_util else 0,
            "complets":                 stats_util["nb_complet"]   if stats_util else 0,
            "incomplets":               stats_util["nb_incomplet"] if stats_util else 0,
            "soumis":                   stats_util["nb_soumis"]    if stats_util else 0,
            "nb_dossiers_avec_relance": nb_dossiers_avec_relance,
            "nb_relances_total":        nb_relances_total,
        },
        "analyse": {
            "nb_dossiers_analyses": len(analyse_rows),
            "risques_top5":        sorted(risques_compteur.items(),     key=lambda x: -x[1])[:5],
            "protections_top5":    sorted(protections_compteur.items(), key=lambda x: -x[1])[:5],
            "acteurs_top10":       sorted(acteurs_compteur.items(),     key=lambda x: -x[1])[:10],
        },
    }


@app.get("/api/v1/dashboard/alertes")
def get_alertes(user=Depends(_get_current_user), db=Depends(_get_db)):
    """Alertes et flags humains pour le Dashboard."""
    flags = db.execute(
        """
        SELECT a.*, d.reference as dossier_reference
        FROM alertes a
        LEFT JOIN dossiers d ON d.id = a.dossier_id
        WHERE a.type_alerte IN ('FLAG_HUMAIN', 'CHAMP_NON_COMMUNIQUE', 'CHAMP_A_COMPLETER_PRO', 'REVUE_INSTRUCTEUR') AND a.acquittee = 0
          -- N'afficher que les alertes de dossiers existants ET non supprimés :
          -- une alerte orpheline ne doit plus faire resurgir un dossier supprimé.
          AND d.id IS NOT NULL AND d.deleted_at IS NULL
        ORDER BY a.created_at DESC LIMIT 50
        """
    ).fetchall()

    relances = db.execute(
        """
        SELECT * FROM relances
        WHERE statut = 'PLANIFIEE'
        ORDER BY planifiee_le ASC LIMIT 50
        """
    ).fetchall()

    return {
        "flags":    [dict(f) for f in flags],
        "relances": [dict(r) for r in relances],
    }


@app.get("/api/v1/dashboard/dossiers/{dossier_id}/appui-evaluation")
def get_appui_evaluation(
    dossier_id: str,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """
    Retourne l'analyse de situation et les données d'appui à l'évaluation.

    Ce panneau éclaire la lecture professionnelle sans orienter automatiquement.
    Il identifie des besoins, des risques, des protections et des acteurs mobilisables.

    IMPORTANT : les acteurs listés sont des pistes, jamais une orientation automatique.
    La décision appartient au professionnel.
    """
    row = db.execute(
        """SELECT analyse_situation_json, synthese_situation, niveau_confiance_analyse,
                  score_maturite_cerfa, niveau_maturite, alertes_qualite_json,
                  rapport_qualite_json, profil_principal, profil_secondaire
           FROM dossiers WHERE id = ?""",
        (dossier_id,),
    ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Dossier non trouvé")

    analyse_dict = json.loads(row["analyse_situation_json"] or "{}")
    rapport_dict = json.loads(row["rapport_qualite_json"]   or "{}")

    return {
        "dossier_id":            dossier_id,
        "niveau_confiance":      row["niveau_confiance_analyse"] or "FAIBLE",
        "score_maturite":        row["score_maturite_cerfa"] or 0,
        "niveau_maturite":       row["niveau_maturite"] or "",
        "profil_principal":      row["profil_principal"] or "",
        "profil_secondaire":     row["profil_secondaire"] or "",
        # Synthèse de situation
        "synthese_situation":    row["synthese_situation"] or "",
        "points_a_completer":    analyse_dict.get("points_a_completer", []),
        # Facteurs
        "facteurs_risque":       analyse_dict.get("facteurs_risque", []),
        "facteurs_protection":   analyse_dict.get("facteurs_protection", []),
        # Besoins
        "besoins_compensation":  analyse_dict.get("besoins_compensation", []),
        "besoins_accompagnement": analyse_dict.get("besoins_accompagnement", []),
        # Vigilances
        "points_vigilance":      analyse_dict.get("points_vigilance", []),
        "alertes_qualite":       json.loads(row["alertes_qualite_json"] or "[]"),
        # Acteurs (pistes, jamais orientations automatiques)
        "acteurs_susceptibles":  analyse_dict.get("acteurs_mobilisables", []),
        # Préconisations
        "preconisations":        analyse_dict.get("preconisations", []),
        # Méta
        "avertissement":         (
            "Ce panneau est un outil d'appui à la décision professionnelle. "
            "Il ne constitue ni une orientation ni une décision. "
            "Toute décision appartient au professionnel accompagnant."
        ),
        # Coverage de collecte (inférenceur)
        "inference_coverage":    json.loads(
            (db.execute("SELECT synthese_json FROM dossiers WHERE id=?", (dossier_id,)).fetchone() or {}).get("synthese_json") or "{}"
        ).get("_inference_coverage") or {},
    }


@app.get("/api/v1/dashboard/dossiers/{dossier_id}/autodetermination-conflicts")
def get_autodetermination_conflicts(
    dossier_id: str,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """
    Retourne l'historique des conflits d'autodétermination pour un dossier.

    Un conflit est levé quand l'usager (WhatsApp) a exprimé un souhait d'orientation
    différent de celui saisi par le professionnel. Le choix de l'usager prime (Règle A).
    Cette alerte permet au professionnel d'en être informé pour adapter son accompagnement.
    """
    row = db.execute(
        "SELECT has_autodetermination_conflict, conflict_history_json FROM dossiers WHERE id = ?",
        (dossier_id,),
    ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Dossier non trouvé")

    return {
        "dossier_id":                    dossier_id,
        "has_autodetermination_conflict": bool(row["has_autodetermination_conflict"]),
        "conflict_history":               json.loads(row["conflict_history_json"] or "[]"),
    }


class CreateDossierRequest(BaseModel):
    departement_code: str
    type_dossier:     str = "INITIAL"


class CreateDossierRequest(BaseModel):
    departement_code:  str
    type_dossier:      str  = "INITIAL"          # INITIAL|RENOUVELLEMENT|SITUATION_CHANGEE|REEVALUATION|AIDANT
    numero_dossier_mdph: str | None = None        # optionnel — ne bloque pas la création
    aidant_familial:   bool = False


@app.post("/api/v1/dashboard/dossiers/create")
def create_dossier_dashboard(
    body: CreateDossierRequest,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """
    Crée un nouveau dossier vide depuis le dashboard professionnel.
    Crée également un usager placeholder si nécessaire (FOREIGN KEY).
    """
    from app.engines.case_state_engine import create_dossier, transition_dossier, DossierStatut
    from app.security.encryption import generate_reference

    # ── Créer un usager placeholder pour ce dossier ───────────────────────────
    # Les dossiers créés manuellement n'ont pas encore de numéro WhatsApp
    usager_id = str(uuid.uuid4())
    ref_usager = generate_reference("FAC")
    now_iso    = datetime.now(timezone.utc).isoformat()
    db.execute(
        """INSERT INTO usagers
            (id, telephone_enc, reference_interne, canal_prefere, created_at, updated_at)
           VALUES (?, '[DASHBOARD]', ?, 'dashboard', ?, ?)""",
        (usager_id, ref_usager, now_iso, now_iso),
    )

    dossier_id = create_dossier(
        usager_id=usager_id,
        departement_code=body.departement_code.strip()[:3],
        organisation_id=user.get("org_id"),
        educateur_id=user.get("sub"),
        type_dossier=body.type_dossier,
        db_conn=db,
    )

    # Stocker le numéro de dossier MDPH existant si renseigné
    if body.numero_dossier_mdph:
        donnees = {"historique_mdph": f"Dossier existant N°{body.numero_dossier_mdph}"}
        db.execute(
            "UPDATE dossiers SET synthese_json = ?, updated_at = ? WHERE id = ?",
            (json.dumps(donnees, ensure_ascii=False), now_iso, dossier_id),
        )

    transition_dossier(
        dossier_id, DossierStatut.BROUILLON, DossierStatut.EN_COLLECTE,
        raison=f"Création manuelle — {body.type_dossier}", canal="dashboard", db_conn=db,
    )
    event_logger.log_event(
        "DOSSIER_CREE_DASHBOARD", dossier_id=dossier_id,
        utilisateur_id=user.get("sub"), canal="dashboard", db_conn=db,
    )
    return {"success": True, "dossier_id": dossier_id}


@app.post("/api/v1/dashboard/dossiers/{dossier_id}/validate-flag")
def validate_flag(
    dossier_id: str,
    request: Request,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """Validation humaine d'un flag — Human-in-the-loop."""
    access_control.require_permission(user.get("role"), access_control.Permission.HUMAN_VALIDATE)
    now = datetime.now(timezone.utc).isoformat()
    db.execute(
        """
        UPDATE dossiers SET
            flag_humain_requis = 0,
            raison_flag        = NULL,
            updated_at         = ?
        WHERE id = ?
        """,
        (now, dossier_id),
    )
    db.execute(
        "UPDATE alertes SET acquittee = 1, acquittee_par = ?, acquittee_le = ? WHERE dossier_id = ?",
        (user.get("sub"), now, dossier_id),
    )
    event_logger.log_event(
        "VALIDATION_HUMAINE",
        dossier_id=dossier_id,
        utilisateur_id=user.get("sub"),
        canal="dashboard",
        db_conn=db,
    )
    return {"success": True}


@app.post("/api/v1/dashboard/dossiers/{dossier_id}/relance")
def send_relance(
    dossier_id: str,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """Envoie une relance manuelle pour un dossier."""
    dossier_row = db.execute("SELECT * FROM dossiers WHERE id = ?", (dossier_id,)).fetchone()
    if not dossier_row:
        raise HTTPException(status_code=404, detail="Dossier non trouvé")

    dossier = dict(dossier_row)
    usager_row = db.execute("SELECT * FROM usagers WHERE id = ?", (dossier.get("usager_id"),)).fetchone()
    if not usager_row:
        raise HTTPException(status_code=400, detail="Usager non trouvé")

    usager = dict(usager_row)
    questions = json.loads(dossier.get("questions_json", "[]") or "[]")

    _notif_service.send_relance_pieces_manquantes(
        usager=usager,
        pieces_manquantes=questions[:5] if questions else ["Informations complémentaires requises"],
        dossier_id=dossier_id,
        db_conn=db,
    )

    relance_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    db.execute(
        """
        INSERT INTO relances
            (id, dossier_id, usager_id, canal, type_relance, planifiee_le, envoyee_le, statut, created_at)
        VALUES (?, ?, ?, 'whatsapp', 'PIECES_MANQUANTES', ?, ?, 'ENVOYEE', ?)
        """,
        (relance_id, dossier_id, usager["id"], now, now, now),
    )
    event_logger.log_event(
        "RELANCE_ENVOYEE",
        dossier_id=dossier_id,
        utilisateur_id=user.get("sub"),
        canal="dashboard",
        db_conn=db,
    )
    return {"success": True}


# ── Pages HTML ────────────────────────────────────────────────────────────────

@app.get("/")
def root():
    from pathlib import Path
    # Résolution absolue pour compatibilité locale ET Railway (Docker WORKDIR=/app)
    base = Path(__file__).resolve().parent.parent
    index_path = base / "static" / "index.html"
    # Fallback : Railway peut poser le projet directement dans /app
    if not index_path.is_file():
        index_path = Path("/app/static/index.html")
    if index_path.is_file():
        return FileResponse(str(index_path), headers={"Cache-Control": "no-cache, no-store, must-revalidate"})
    logger.warning("[root] index.html introuvable — chemins tentés : %s", index_path)
    return RedirectResponse("/dashboard")


@app.get("/dashboard")
def dashboard_page(request: Request):
    """Sert le dashboard. Vérifie le cookie JWT côté serveur avant d'envoyer le HTML."""
    token = request.cookies.get("facilim_token")
    if not token:
        return RedirectResponse("/login", status_code=302)
    payload = token_manager.verify_dashboard_token(token, settings.jwt_secret_key)
    if not payload:
        return RedirectResponse("/login", status_code=302)
    dashboard_path = os.path.join(os.path.dirname(__file__), "dashboards", "dashboard.html")
    if os.path.isfile(dashboard_path):
        return FileResponse(dashboard_path, headers={"Cache-Control": "no-cache, no-store, must-revalidate"})
    return HTMLResponse("<h1>Facilim v2 — Dashboard en cours de chargement</h1>")


@app.get("/login")
def login_page():
    static_login = os.path.join(os.path.dirname(__file__), "..", "static", "login.html")
    if os.path.isfile(static_login):
        return FileResponse(static_login, headers={"Cache-Control": "no-cache, no-store, must-revalidate"})
    return HTMLResponse("""
    <!DOCTYPE html><html lang="fr"><head><meta charset="UTF-8">
    <title>Facilim — Connexion</title>
    <script src="https://cdn.tailwindcss.com"></script></head>
    <body class="bg-gray-50 flex items-center justify-center min-h-screen">
    <div class="bg-white p-8 rounded-xl shadow w-full max-w-sm">
      <h1 class="text-xl font-bold text-gray-900 mb-6">Facilim — Connexion</h1>
      <form id="f" class="space-y-4">
        <input id="email" type="email" placeholder="Email" required
          class="w-full border rounded-lg px-3 py-2 text-sm">
        <input id="pass" type="password" placeholder="Mot de passe" required
          class="w-full border rounded-lg px-3 py-2 text-sm">
        <button type="submit" class="w-full bg-blue-700 text-white py-2 rounded-lg text-sm font-medium">
          Se connecter
        </button>
        <p id="err" class="text-red-600 text-xs hidden">Identifiants incorrects</p>
      </form>
      <script>
        document.getElementById('f').addEventListener('submit', async e => {
          e.preventDefault();
          const res = await fetch('/api/v1/auth/login', {
            method:'POST', credentials:'include',
            headers:{'Content-Type':'application/json'},
            body: JSON.stringify({email: document.getElementById('email').value,
                                  password: document.getElementById('pass').value})
          });
          if (res.ok) window.location='/dashboard';
          else document.getElementById('err').classList.remove('hidden');
        });
      </script>
    </div></body></html>
    """)


# ── RGPD ─────────────────────────────────────────────────────────────────────

@app.get("/api/v1/rgpd/export/{usager_id}")
def export_rgpd(
    usager_id: str,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """Export des données d'un usager (Article 15 RGPD)."""
    access_control.require_permission(user.get("role"), access_control.Permission.RGPD_MANAGE)
    from app.services.compliance_service import ComplianceService
    svc = ComplianceService(retention_days=settings.rgpd_retention_days)
    data = svc.export_usager_data(
        usager_id=usager_id,
        db_conn=db,
        requesting_utilisateur_id=user.get("sub"),
    )
    return JSONResponse(data)


@app.delete("/api/v1/rgpd/erase/{usager_id}")
def erase_rgpd(
    usager_id: str,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """Effacement des données d'un usager (Article 17 RGPD)."""
    access_control.require_permission(user.get("role"), access_control.Permission.RGPD_MANAGE)
    from app.services.compliance_service import ComplianceService
    svc = ComplianceService(retention_days=settings.rgpd_retention_days)
    ok = svc.erase_usager_data(
        usager_id=usager_id,
        db_conn=db,
        requesting_utilisateur_id=user.get("sub"),
    )
    return {"success": ok}


# ── Audit ─────────────────────────────────────────────────────────────────────

@app.get("/api/v1/audit/events")
def get_audit_events(
    dossier_id: str | None = None,
    limit: int = 100,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """Lecture du journal d'audit (SUPER_ADMIN / ADMIN_ESSMS uniquement)."""
    access_control.require_permission(user.get("role"), access_control.Permission.AUDIT_READ)
    if dossier_id:
        rows = db.execute(
            "SELECT * FROM audit_events WHERE dossier_id = ? ORDER BY created_at DESC LIMIT ?",
            (dossier_id, limit),
        ).fetchall()
    else:
        rows = db.execute(
            "SELECT * FROM audit_events ORDER BY created_at DESC LIMIT ?",
            (limit,),
        ).fetchall()
    return {"events": [dict(r) for r in rows]}


@app.get("/api/v1/dossiers/{dossier_id}/historique")
def get_historique_dossier(
    dossier_id: str,
    limit: int = 200,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """
    Historique complet et chronologique d'un dossier (traçabilité PROD-4).

    Fusionne en LECTURE SEULE les deux journaux DÉJÀ persistés — `audit_events`
    (event_logger) et `cerfa_audit_log` (_log_cerfa_audit) — sans rien recalculer
    ni appeler aucun moteur. Permet de reconstruire « qui a fait quoi, quand, sur
    quel dossier, avec quel résultat ».
    """
    evenements: list[dict] = []

    # 1) Journal général — audit_events
    try:
        for r in db.execute(
            "SELECT type_evenement, created_at, canal, utilisateur_id, payload_json "
            "FROM audit_events WHERE dossier_id = ? ORDER BY created_at",
            (dossier_id,),
        ).fetchall():
            d = dict(r)
            evenements.append({
                "source":  "audit_events",
                "type":    d.get("type_evenement"),
                "date":    d.get("created_at"),
                "canal":   d.get("canal"),
                "acteur":  d.get("utilisateur_id"),
                "details": json.loads(d.get("payload_json") or "{}"),
            })
    except Exception as e:
        logger.warning("[HISTORIQUE] audit_events indisponible : %s", e)

    # 2) Journal CERFA — cerfa_audit_log (cockpit absent ici ; gate/export y figurent)
    try:
        for r in db.execute(
            "SELECT event_type, event_at, canal, acteur_id, acteur_type, details_json "
            "FROM cerfa_audit_log WHERE dossier_id = ? ORDER BY event_at",
            (dossier_id,),
        ).fetchall():
            d = dict(r)
            evenements.append({
                "source":  "cerfa_audit_log",
                "type":    d.get("event_type"),
                "date":    d.get("event_at"),
                "canal":   d.get("canal"),
                "acteur":  d.get("acteur_id"),
                "details": json.loads(d.get("details_json") or "{}"),
            })
    except Exception as e:
        logger.warning("[HISTORIQUE] cerfa_audit_log indisponible : %s", e)

    # Tri chronologique croissant (ordre des événements)
    evenements.sort(key=lambda ev: ev.get("date") or "")
    return {
        "dossier_id":    dossier_id,
        "nb_evenements": len(evenements),
        "evenements":    evenements[:limit],
    }


# ── Compatibilité avec l'ancien main.py (route legacy) ───────────────────────
# ── Analytics V2.2 — Mesure · Analyse · Pilotage ─────────────────────────────
# Module totalement indépendant du pipeline CERFA.
# Aucune modification du moteur métier.

class EvaluationInput(BaseModel):
    """Données saisies par le professionnel après correction d'un dossier."""
    dossier_id:               str | None = None
    date_evaluation:          str        = ""
    evaluateur:               str        = "nassim"
    version_facilim:          str        = ""
    essms_source:             str        = ""
    profil_mdph:              str        = ""
    type_handicap:            str        = ""
    type_dossier:             str        = ""
    document_present:         bool       = False
    type_document:            str        = ""
    niveau_documentation:     int        = 0      # 0–3
    score_b:                  float      = -1
    score_c:                  float      = -1
    score_d:                  float      = -1
    score_e:                  float      = -1
    score_maturite:           int        = 0
    temps_correction:         float      = 0
    temps_correction_b:       float      = 0
    temps_correction_c:       float      = 0
    temps_correction_d:       float      = 0
    temps_correction_e:       float      = 0
    temps_estime_sans_facilim: float     = 0
    temps_reel_avec_facilim:  float      = 0
    taux_reecriture_b:        float      = 0
    taux_reecriture_c:        float      = 0
    taux_reecriture_d:        float      = 0
    taux_reecriture_e:        float      = 0
    questions_supprimees:     int        = 0
    items_extraits:           int        = 0
    erreurs_detectees:        list[str]  = []
    commentaire:              str        = ""


@app.post("/api/v1/analytics/evaluations")
def post_evaluation(
    body: EvaluationInput,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """
    Enregistre une évaluation terrain.
    Calcule automatiquement score_global, gain_temps et niveau_validation.
    """
    from app.analytics.quality_dashboard import enregistrer_evaluation
    eval_id = enregistrer_evaluation(db, body.model_dump())
    return {"success": True, "evaluation_id": eval_id}


@app.get("/api/v1/analytics/dashboard")
def get_analytics_dashboard(
    debut: str = "",
    fin:   str = "",
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """
    Tableau de bord JSON complet — tous les KPI analytiques.
    Paramètres optionnels : debut=YYYY-MM-DD, fin=YYYY-MM-DD
    """
    from app.analytics.quality_dashboard import generer_tableau_bord
    try:
        return generer_tableau_bord(db, periode_debut=debut, periode_fin=fin)
    except Exception as e:
        logger.error("[ANALYTICS] Dashboard error : %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/v1/analytics/rapport")
def get_analytics_rapport(
    debut: str = "",
    fin:   str = "",
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """Rapport exécutif texte — réponse aux 3 questions stratégiques."""
    from app.analytics.quality_dashboard import generer_rapport_executif
    try:
        texte = generer_rapport_executif(db, periode_debut=debut, periode_fin=fin)
        return {"rapport": texte}
    except Exception as e:
        logger.error("[ANALYTICS] Rapport error : %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# Conservée pour ne pas casser les intégrations existantes

@app.get("/api/v1/dashboard/nouveau-dossier")
def redirect_nouveau_dossier():
    """Redirige l'ancienne URL vers le dashboard (cache busting)."""
    return RedirectResponse("/dashboard", status_code=302)


@app.get("/api/v1/dossiers")
def legacy_list_dossiers(user=Depends(_get_current_user), db=Depends(_get_db)):
    """Route legacy — compatible avec static/dashboard.html."""
    rows = db.execute(
        "SELECT * FROM dossiers WHERE deleted_at IS NULL ORDER BY updated_at DESC LIMIT 100"
    ).fetchall()
    dossiers = []
    for row in rows:
        d = dict(row)
        d["droits_identifies"] = json.loads(d.get("droits_identifies_json", "[]") or "[]")
        d["synthese"] = json.loads(d.get("synthese_json", "{}") or "{}")
        d["questions"] = json.loads(d.get("questions_json", "[]") or "[]")
        dossiers.append(d)
    return {"dossiers": dossiers}


# ── Endpoints legacy (static/dashboard.html) ─────────────────────────────────

@app.post("/api/v1/dossiers/{dossier_id}/force-relance")
def force_relance(
    dossier_id: str,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """Force une relance WhatsApp immédiate."""
    dossier_row = db.execute("SELECT * FROM dossiers WHERE id = ?", (dossier_id,)).fetchone()
    if not dossier_row:
        raise HTTPException(status_code=404, detail="Dossier non trouvé")
    dossier = dict(dossier_row)
    usager = db.execute("SELECT * FROM usagers WHERE id = ?", (dossier.get("usager_id"),)).fetchone()
    if not usager:
        raise HTTPException(status_code=400, detail="Usager non trouvé")
    questions = json.loads(dossier.get("questions_json", "[]") or "[]")
    _notif_service.send_relance_pieces_manquantes(
        usager=dict(usager),
        pieces_manquantes=questions[:5] if questions else ["Informations complémentaires requises"],
        dossier_id=dossier_id,
        db_conn=db,
    )
    event_logger.log_event("RELANCE_FORCEE", dossier_id=dossier_id,
                           utilisateur_id=user.get("sub"), canal="dashboard", db_conn=db)
    return {"success": True}


class DossierInitiateRequest(BaseModel):
    telephone:               str
    email:                   str          # obligatoire — nécessaire pour envoyer le CERFA final
    departement_code:        str         = "75"
    type_dossier:            str         = "INITIAL"
    langue:                  str         = "fr"
    urgent:                  bool        = False
    nom_prenom:              str | None  = None
    date_naissance:          str | None  = None
    notes_pro:               str | None  = None
    numero_dossier_mdph:     str | None  = None
    numero_securite_sociale: str | None  = None
    aidant_familial:         bool        = False


@app.post("/api/v1/dossiers/initiate")
def initiate_dossier(
    body: DossierInitiateRequest,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """
    Initie un dossier depuis le dashboard pro.
    - Crée ou retrouve l'usager par téléphone
    - Pré-remplit le dossier avec les données connues
    - Envoie automatiquement le message WhatsApp de démarrage
    """
    from app.engines.case_state_engine import create_dossier, transition_dossier, DossierStatut
    from app.security.encryption import encrypt, generate_reference

    if not body.telephone.strip():
        raise HTTPException(status_code=400, detail="Le numéro WhatsApp est obligatoire.")
    if not body.email.strip():
        raise HTTPException(status_code=400, detail="L'email est obligatoire (envoi du CERFA final).")

    now_iso   = datetime.now(timezone.utc).isoformat()
    # Normaliser en format LOCAL canonique avant chiffrement (0XXXXXXXXX)
    # Garantit que encrypt("0642...") == encrypt("0642...") quelle que soit la saisie du pro
    _raw = body.telephone.strip().lstrip("+")
    phone_raw = ("0" + _raw[2:]) if (_raw.startswith("33") and len(_raw) == 11) else _raw
    phone_enc = encrypt(phone_raw)

    # Recherche ou création de l'usager
    usager_row = db.execute(
        "SELECT * FROM usagers WHERE telephone_enc = ? LIMIT 1", (phone_enc,)
    ).fetchone()

    if usager_row:
        usager_id = usager_row["id"]
        # Mettre à jour l'email si nouveau
        if body.email:
            db.execute(
                "UPDATE usagers SET email_enc = ?, updated_at = ? WHERE id = ?",
                (encrypt(body.email), now_iso, usager_id),
            )
    else:
        usager_id = str(uuid.uuid4())
        ref       = generate_reference("FAC")
        email_enc = encrypt(body.email)
        db.execute(
            """INSERT INTO usagers
               (id, telephone_enc, email_enc, reference_interne, canal_prefere, langue_preferee, created_at, updated_at)
               VALUES (?, ?, ?, ?, 'whatsapp', ?, ?, ?)""",
            (usager_id, phone_enc, email_enc, ref, body.langue, now_iso, now_iso),
        )

    # Création du dossier
    dossier_id = create_dossier(
        usager_id=usager_id,
        departement_code=body.departement_code.strip()[:3],
        educateur_id=user.get("sub"),
        type_dossier=body.type_dossier,
        db_conn=db,
    )

    # Pré-remplissage synthese_json avec TOUTES les données connues → CERFA automatique
    donnees_initiales: dict = {
        "telephone":        phone_raw,
        "email":            body.email,
        "departement":      body.departement_code.strip()[:3],
        "type_dossier":     body.type_dossier,
        "langue":           body.langue,
    }
    if body.nom_prenom:           donnees_initiales["nom_prenom"]          = body.nom_prenom
    # Ne pas stocker un gabarit de date (« JJ/MM/AAAA ») : seule une vraie date,
    # contenant des chiffres, est conservée — sinon le champ reste à collecter.
    if body.date_naissance and any(c.isdigit() for c in str(body.date_naissance)):
        donnees_initiales["date_naissance"] = body.date_naissance
    if body.notes_pro:            donnees_initiales["notes_pro"]           = body.notes_pro
    if body.numero_securite_sociale:
        donnees_initiales["numero_securite_sociale"] = body.numero_securite_sociale.replace(" ", "")
        donnees_initiales["num_secu"]                = body.numero_securite_sociale.replace(" ", "")
    if body.numero_dossier_mdph:
        donnees_initiales["historique_mdph"]     = f"Dossier existant N°{body.numero_dossier_mdph}"
        donnees_initiales["numero_dossier_mdph"] = body.numero_dossier_mdph
    if body.aidant_familial:      donnees_initiales["aidant_demande"]      = True
    if body.urgent:               donnees_initiales["urgence"]             = True

    from app.engines.orchestration_engine import _score_completude_metier
    score_init = _score_completude_metier(donnees_initiales)   # P4 : complétude métier
    db.execute(
        "UPDATE dossiers SET synthese_json = ?, score_completude = ?, updated_at = ? WHERE id = ?",
        (json.dumps(donnees_initiales, ensure_ascii=False), score_init, now_iso, dossier_id),
    )

    # Création de la session WhatsApp pour le suivi conversationnel
    session_id = str(uuid.uuid4())
    db.execute(
        """INSERT INTO sessions_whatsapp
           (id, usager_id, telephone, dossier_id, persona_actif, derniere_activite, created_at)
           VALUES (?, ?, ?, ?, 'identification', ?, ?)""",
        (session_id, usager_id, phone_raw, dossier_id, now_iso, now_iso),
    )

    transition_dossier(
        dossier_id, DossierStatut.BROUILLON, DossierStatut.EN_COLLECTE,
        raison=f"Initié depuis dashboard — {body.type_dossier}", canal="dashboard", db_conn=db,
    )
    db.commit()  # commit avant d'envoyer le message

    # Le professionnel déclenche la conversation manuellement via le bouton "Lancer"
    # Ici on NE ENVOIE PAS encore le message WhatsApp — le pro doit cliquer "Lancer la conversation"

    event_logger.log_event(
        "DOSSIER_INITIE_DASHBOARD", dossier_id=dossier_id,
        utilisateur_id=user.get("sub"), canal="dashboard", db_conn=db,
    )
    return {"success": True, "dossier_id": dossier_id}


class AjouterTelephoneRequest(BaseModel):
    telephone: str
    email:     str | None = None


@app.post("/api/v1/dossiers/{dossier_id}/ajouter-telephone")
def ajouter_telephone_dossier(
    dossier_id: str,
    body: AjouterTelephoneRequest,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """
    Ajoute ou met à jour le téléphone WhatsApp d'un dossier existant.
    Crée la session WhatsApp si elle n'existe pas.
    """
    from app.security.encryption import encrypt, generate_reference
    dossier_row = db.execute("SELECT * FROM dossiers WHERE id = ?", (dossier_id,)).fetchone()
    if not dossier_row:
        raise HTTPException(status_code=404, detail="Dossier non trouvé")

    phone     = body.telephone.strip()
    now       = datetime.now(timezone.utc).isoformat()
    phone_enc = encrypt(phone)

    # Trouver ou créer l'usager
    usager_row = db.execute(
        "SELECT id FROM usagers WHERE telephone_enc = ? LIMIT 1", (phone_enc,)
    ).fetchone()
    if usager_row:
        usager_id = usager_row["id"]
    else:
        usager_id = str(uuid.uuid4())
        ref = generate_reference("FAC")
        email_enc = encrypt(body.email) if body.email else None
        db.execute(
            "INSERT INTO usagers (id, telephone_enc, email_enc, reference_interne, canal_prefere, created_at, updated_at) VALUES (?, ?, ?, ?, 'whatsapp', ?, ?)",
            (usager_id, phone_enc, email_enc, ref, now, now),
        )

    # Mettre à jour le dossier avec l'usager
    db.execute(
        "UPDATE dossiers SET usager_id = ?, updated_at = ? WHERE id = ?",
        (usager_id, now, dossier_id),
    )

    # Créer ou mettre à jour la session WhatsApp
    existing_session = db.execute(
        "SELECT id FROM sessions_whatsapp WHERE dossier_id = ? LIMIT 1", (dossier_id,)
    ).fetchone()
    if existing_session:
        db.execute(
            "UPDATE sessions_whatsapp SET telephone = ?, derniere_activite = ? WHERE id = ?",
            (phone, now, existing_session["id"]),
        )
    else:
        session_id = str(uuid.uuid4())
        db.execute(
            "INSERT INTO sessions_whatsapp (id, usager_id, telephone, dossier_id, persona_actif, derniere_activite, created_at) VALUES (?, ?, ?, ?, 'identification', ?, ?)",
            (session_id, usager_id, phone, dossier_id, now, now),
        )
    db.commit()
    return {"success": True, "telephone": phone[-4:]}


@app.post("/api/v1/dossiers/{dossier_id}/lancer-conversation")
def lancer_conversation_whatsapp(
    dossier_id: str,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """
    Le professionnel déclenche manuellement la conversation WhatsApp avec l'usager.
    Appelé après avoir alimenté le dossier (documents, notes).
    Envoie le premier message d'accueil avec les informations déjà connues.
    """
    dossier_row = db.execute("SELECT * FROM dossiers WHERE id = ?", (dossier_id,)).fetchone()
    if not dossier_row:
        raise HTTPException(status_code=404, detail="Dossier non trouvé")
    dossier = dict(dossier_row)  # convertir sqlite3.Row en dict

    # Récupérer le numéro WhatsApp depuis la session
    session = db.execute(
        "SELECT telephone FROM sessions_whatsapp WHERE dossier_id = ? LIMIT 1",
        (dossier_id,)
    ).fetchone()
    if not session or not session["telephone"]:
        raise HTTPException(status_code=400, detail="Aucun numéro WhatsApp associé à ce dossier. Veuillez ajouter le numéro de téléphone.")

    # Normaliser en format international (33XXXXXXXXX) requis par l'API WhatsApp Cloud
    _phone_raw = session["telephone"].strip().lstrip("+")
    phone = ("33" + _phone_raw[1:]) if (_phone_raw.startswith("0") and len(_phone_raw) == 10) else _phone_raw
    donnees  = json.loads(dossier.get("synthese_json", "{}") or "{}")
    prenom   = (donnees.get("nom_prenom", "") or "").split()[0] if donnees.get("nom_prenom") else ""
    civilite = f"Bonjour{' ' + prenom if prenom else ''} !\n\n"

    # Compter les infos déjà connues grâce aux documents uploadés
    infos_connues = [k for k, v in donnees.items()
                     if v and not k.startswith("_") and k not in ("notes_pro", "email")]
    nb_connus = len(infos_connues)

    # Le premier message est UNIQUEMENT la demande de consentement RGPD.
    # Aucune question sur le dossier à ce stade.
    msg = (
        f"{civilite}"
        f"Je suis de l'équipe Facilim. Votre professionnel accompagnant a ouvert un dossier MDPH pour vous.\n\n"
        f"Avant tout, j'ai besoin de votre accord pour utiliser vos informations dans le cadre de votre demande MDPH.\n\n"
        f"📋 *Vos données serviront uniquement à :*\n"
        f"• Préparer votre dossier MDPH\n"
        f"• Partager votre dossier avec la MDPH de votre département\n"
        f"• Vous tenir informé(e) de l'avancement\n\n"
        f"🔒 Vos données médicales sont protégées et chiffrées. Elles ne seront jamais vendues.\n\n"
        f"Vous pouvez demander à voir, corriger ou effacer vos données à tout moment.\n"
        f"📧 Contact : contactfacilim@gmail.com\n\n"
        f"*Tapez OUI pour accepter ou NON pour refuser.*"
    )

    try:
        _wa_service.send_text(phone, msg, dossier_id=dossier_id, db_conn=db)
        now = datetime.now(timezone.utc).isoformat()
        db.execute("UPDATE dossiers SET updated_at = ? WHERE id = ?", (now, dossier_id))
        # Commit immédiat : garantit que synthese_json est visible par le pipeline WhatsApp
        db.commit()
        event_logger.log_event("CONVERSATION_LANCEE", dossier_id=dossier_id,
                               utilisateur_id=user.get("sub"), canal="dashboard", db_conn=db)
        return {"success": True, "phone": phone[-4:], "nb_infos_connues": nb_connus}
    except Exception as e:
        logger.error("[LANCER] Erreur envoi WhatsApp : %s", e)
        raise HTTPException(status_code=500, detail=f"Envoi WhatsApp échoué : {e}")


@app.post("/api/v1/extract-text")
async def extract_text(
    file: UploadFile = File(...),
    user=Depends(_get_current_user),
):
    """Extrait le texte d'un PDF ou d'une image uploadée (sans lien à un dossier)."""
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Fichier vide")
    texte = _extraire_texte_fichier(content, file.filename or "", file.content_type or "")
    return {"texte": texte, "nb_chars": len(texte), "filename": file.filename}


def _extraire_texte_fichier(content: bytes, filename: str, content_type: str) -> str:
    """
    Extrait le texte brut depuis :
     - PDF  → pypdf
     - Word (.docx) → python-docx (paragraphes + tableaux)
     - Image → message informatif (OCR non disponible)
    """
    import io
    fn_lower = filename.lower()
    try:
        # ── PDF ──────────────────────────────────────────────────────────────
        if "pdf" in content_type or fn_lower.endswith(".pdf"):
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)

        # ── Word .docx ────────────────────────────────────────────────────────
        elif fn_lower.endswith(".docx") or "wordprocessingml" in content_type:
            from docx import Document
            doc   = Document(io.BytesIO(content))
            lines = []
            for para in doc.paragraphs:
                t = para.text.strip()
                if t:
                    lines.append(t)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        lines.append(" | ".join(cells))
            return "\n".join(lines)

        # ── Excel .xlsx ───────────────────────────────────────────────────────
        elif fn_lower.endswith((".xlsx", ".xls")) or "spreadsheet" in content_type or "excel" in content_type:
            import openpyxl
            wb    = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            lines = []
            for ws in wb.worksheets:
                lines.append(f"=== Feuille : {ws.title} ===")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                    if cells:
                        lines.append(" | ".join(cells))
            return "\n".join(lines)

        # ── Texte brut ────────────────────────────────────────────────────────
        elif fn_lower.endswith((".txt", ".csv")) or "text/plain" in content_type:
            try:
                return content.decode("utf-8", errors="replace")
            except Exception:
                return content.decode("latin-1", errors="replace")

        # ── Image ─────────────────────────────────────────────────────────────
        elif any(fn_lower.endswith(ext) for ext in (".jpg", ".jpeg", ".png", ".webp", ".gif")):
            return f"[Image reçue : {filename} — enregistrée pour validation humaine]"

        else:
            return f"[Format non pris en charge : {filename}]"

    except Exception as e:
        return f"[Extraction échouée pour {filename} : {e}]"


@app.post("/api/v1/dossiers/{dossier_id}/upload-document")
async def upload_document_dossier(
    dossier_id: str,
    file: UploadFile = File(...),
    notes: str = Form(default=""),
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """
    Upload un document lié à un dossier.
    Extrait automatiquement le texte/données et enrichit synthese_json.
    Le pro peut uploader un certificat médical, un compte-rendu, etc.
    """
    dossier_row = db.execute("SELECT * FROM dossiers WHERE id = ?", (dossier_id,)).fetchone()
    if not dossier_row:
        raise HTTPException(status_code=404, detail="Dossier non trouvé")

    content   = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Fichier vide")

    filename     = file.filename or "document"
    content_type = file.content_type or ""
    texte_extrait = _extraire_texte_fichier(content, filename, content_type)

    # Enrichissement des données via LLM si texte significatif
    donnees_extraites: dict = {}
    if len(texte_extrait) > 50 and not texte_extrait.startswith("["):
        try:
            import json as _json
            prompt = f"""Tu es un assistant spécialisé MDPH. Analyse ce document (bilan médical, compte-rendu, bilan de parcours professionnel, certificat médical, etc.) et extrais toutes les informations utiles pour un dossier MDPH.

Retourne UNIQUEMENT un JSON avec les champs trouvés parmi cette liste :
- nom_prenom : NOM Prénom (tel qu'écrit dans le document)
- date_naissance : JJ/MM/AAAA
- genre : "homme" ou "femme", déduit de la civilité écrite dans le document (« Mr », « M. », « Monsieur » → "homme" ; « Mme », « Madame », « Mlle » → "femme"). N'invente rien : si aucune civilité, n'inclus pas ce champ.
- commune_naissance : commune de naissance si mentionnée
- pays_naissance : pays de naissance si mentionné
- adresse_complete : adresse postale complète
- departement : code département (2 ou 3 chiffres)
- num_secu : numéro de sécurité sociale 15 chiffres
- email : adresse email si présente dans le document
- telephone : numéro de téléphone si présent
- numero_dossier_mdph : numéro de dossier MDPH existant si mentionné
- diagnostics : liste des diagnostics/pathologies séparés par virgules
- traitements : médicaments ou thérapies en cours
- medecin_traitant : nom du médecin traitant
- impact_quotidien : description des limitations fonctionnelles (ce que la personne ne peut pas faire)
- statut_emploi : situation professionnelle actuelle
- historique_mdph : informations sur les droits MDPH précédents
- accident_travail : description si accident de travail mentionné
- restrictions_emploi : limitations pour le travail (port de charges, position, etc.)
- representant_legal_nom : nom du représentant légal si mentionné
- representant_legal_lien : lien avec la personne (parent, tuteur, curateur…)

Ne retourne QUE les champs que tu trouves dans le document. Si un champ est absent, ne l'inclus pas.

Document :
{texte_extrait[:4000]}
"""
            resp = _llm_client.chat.completions.create(
                model=settings.openai_model_fast,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=400,
                temperature=0.0,
                response_format={"type": "json_object"},
            )
            donnees_extraites = _json.loads(resp.choices[0].message.content)
        except Exception as e:
            logger.warning("[UPLOAD] Extraction LLM échouée : %s", e)

    # Fusionner dans synthese_json
    synthese = json.loads(dossier_row["synthese_json"] or "{}")
    for k, v in donnees_extraites.items():
        if v and k not in synthese:  # ne pas écraser les données existantes
            synthese[k] = v

    # Ajouter le texte brut du document dans les notes pro
    if texte_extrait and not texte_extrait.startswith("["):
        existing_notes = synthese.get("documents_texte", "")
        synthese["documents_texte"] = f"{existing_notes}\n\n--- {filename} ---\n{texte_extrait[:2000]}"

    if notes:
        existing = synthese.get("notes_pro", "")
        synthese["notes_pro"] = f"{existing}\n{notes}".strip()

    now = datetime.now(timezone.utc).isoformat()
    db.execute(
        "UPDATE dossiers SET synthese_json = ?, updated_at = ? WHERE id = ?",
        (json.dumps(synthese, ensure_ascii=False), now, dossier_id),
    )

    # Enregistrer comme pièce justificative
    piece_id = str(uuid.uuid4())
    db.execute(
        """INSERT INTO pieces_justificatives
           (id, dossier_id, type_piece, nom_fichier_original, mime_type, taille_octets,
            ocr_effectue, uploaded_par, created_at, updated_at)
           VALUES (?, ?, 'DOCUMENT_PRO', ?, ?, ?, 1, 'dashboard', ?, ?)""",
        (piece_id, dossier_id, filename, content_type, len(content), now, now),
    )
    event_logger.log_event("DOCUMENT_UPLOADE_PRO", dossier_id=dossier_id,
                           utilisateur_id=user.get("sub"), canal="dashboard",
                           payload={"filename": filename, "champs_extraits": list(donnees_extraites.keys())},
                           db_conn=db)

    return {
        "success":         True,
        "piece_id":        piece_id,
        "champs_extraits": list(donnees_extraites.keys()),
        "nb_chars_texte":  len(texte_extrait),
        "message":         f"{len(donnees_extraites)} champ(s) pré-rempli(s) depuis le document.",
    }


@app.post("/api/v1/dossiers/{dossier_id}/valider-droits")
def valider_droits(
    dossier_id: str,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """Valide les droits identifiés et prépare le CERFA."""
    dossier_row = db.execute("SELECT * FROM dossiers WHERE id = ?", (dossier_id,)).fetchone()
    if not dossier_row:
        raise HTTPException(status_code=404, detail="Dossier non trouvé")
    access_control.require_permission(user.get("role"), access_control.Permission.HUMAN_VALIDATE)
    now = datetime.now(timezone.utc).isoformat()
    db.execute(
        "UPDATE dossiers SET statut = 'COMPLET', flag_humain_requis = 0, updated_at = ? WHERE id = ?",
        (now, dossier_id),
    )
    event_logger.log_event("DROITS_VALIDES", dossier_id=dossier_id,
                           utilisateur_id=user.get("sub"), canal="dashboard", db_conn=db)
    return {"success": True, "statut": "COMPLET"}


# ── Fonctions utilitaires de traçabilité CERFA — Point 3 ─────────────────────

def _anonymiser_email_audit(email: str) -> dict:
    """
    Minimisation RGPD : retourne hash SHA256 + domaine uniquement.
    L'email complet n'est jamais stocké dans cerfa_audit_log.

    Exemples :
      "nassim@facilim.fr"  → {"email_hash": "abc...", "email_domaine": "@facilim.fr"}
      ""                   → {"email_hash": "",         "email_domaine": ""}
    """
    import hashlib
    email = (email or "").strip().lower()
    if not email:
        return {"email_hash": "", "email_domaine": ""}
    h = hashlib.sha256(email.encode("utf-8")).hexdigest()
    domaine = "@" + email.split("@")[-1] if "@" in email else "@inconnu"
    return {"email_hash": h, "email_domaine": domaine}


def _log_cerfa_audit(
    db,
    dossier_id: str,
    event_type: str,
    canal: str | None,
    acteur_id: str | None,
    acteur_type: str,
    details: dict,
) -> None:
    """
    Insère une ligne dans cerfa_audit_log.
    Silencieux en cas d'erreur : le journal ne doit jamais bloquer le flux principal.
    """
    import uuid as _uuid
    import json as _json
    try:
        now = datetime.now(timezone.utc).isoformat()
        db.execute(
            """INSERT INTO cerfa_audit_log
               (id, dossier_id, event_type, event_at, canal, acteur_id, acteur_type,
                details_json, created_at)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                str(_uuid.uuid4()),
                dossier_id,
                event_type,
                now,
                canal,
                acteur_id,
                acteur_type,
                _json.dumps(details, ensure_ascii=False) if details else None,
                now,
            ),
        )
    except Exception as e:
        logger.warning("[AUDIT_CERFA] Journalisation ignorée (%s) — %s", event_type, e)


# ── Constantes Point 2 ───────────────────────────────────────────────────────

_TEXTE_VALIDATION_PROFESSIONNEL = (
    "Je confirme avoir vérifié le dossier, l'avoir remis au bénéficiaire "
    "ou à son représentant légal pour relecture, et l'avoir informé que "
    "la validation finale lui appartient."
)


# ── Routes Point 2 ────────────────────────────────────────────────────────────

@app.post("/api/v1/dossiers/{dossier_id}/valider-transmission")
def valider_transmission(
    dossier_id: str,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """
    Validation professionnelle avant envoi.
    Vérifie les alertes bloquantes (B1-B3) du comité d'agents.
    N'exige PAS encore la validation usager ici — vérifiée à l'envoi.
    INSERT dans cerfa_validations (type=professionnel).
    """
    import hashlib as _hl
    from app.engines.comite_agents import executer_comite_agents

    dossier_row = db.execute("SELECT * FROM dossiers WHERE id = ?", (dossier_id,)).fetchone()
    if not dossier_row:
        raise HTTPException(status_code=404, detail="Dossier non trouvé")

    donnees = json.loads(dossier_row.get("synthese_json", "{}") or "{}")

    # Exécuter le comité en mode pré-validation
    session_row = db.execute(
        "SELECT persona_actif FROM sessions_whatsapp WHERE dossier_id = ? LIMIT 1",
        (dossier_id,)
    ).fetchone()
    profil = (session_row["persona_actif"] if session_row else "adulte") or "adulte"

    rapport = executer_comite_agents(
        donnees=donnees, profil=profil,
        dossier_id=dossier_id, db=db, mode="pre_validation"
    )

    # Journaliser le rapport
    _log_cerfa_audit(
        db=db, dossier_id=dossier_id,
        event_type="analyse_agents", canal="dashboard",
        acteur_id=user.get("sub"), acteur_type="professionnel",
        details={
            "score":              rapport.score,
            "alertes_bloquantes": [a.message for a in rapport.alertes_bloquantes],
            "alertes_mineures":   [a.message for a in rapport.alertes_mineures],
            "pieces_manquantes":  rapport.pieces_manquantes,
            "droits_potentiels":  rapport.droits_potentiels,
            "agents_actives":     rapport.agents_actives,
            "bloquant":           rapport.bloquant,
        },
    )

    if rapport.bloquant:
        raise HTTPException(
            status_code=400,
            detail={
                "message": "Des alertes bloquantes empêchent la validation.",
                "alertes_bloquantes": [a.message for a in rapport.alertes_bloquantes],
                "score": rapport.score,
            }
        )

    # INSERT cerfa_validations
    import uuid as _uuid
    now = datetime.now(timezone.utc).isoformat()
    synthese_hash = _hl.sha256(
        json.dumps(donnees, sort_keys=True, ensure_ascii=False).encode("utf-8")
    ).hexdigest()
    version = dossier_row.get("version") or 1

    db.execute(
        """INSERT INTO cerfa_validations
           (id, dossier_id, validated_by, canal, type_validation, validated_at,
            synthese_hash, cerfa_pdf_hash, dossier_version, confirmation_texte,
            reponse_usager, created_at)
           VALUES (?,?,?,?,?,?,?,?,?,?,?,?)""",
        (
            str(_uuid.uuid4()), dossier_id, user.get("sub"), "dashboard",
            "professionnel", now, synthese_hash, None, version,
            _TEXTE_VALIDATION_PROFESSIONNEL, "OUI", now,
        )
    )
    _log_cerfa_audit(
        db=db, dossier_id=dossier_id,
        event_type="validation_professionnel", canal="dashboard",
        acteur_id=user.get("sub"), acteur_type="professionnel",
        details={"synthese_hash": synthese_hash, "version": version},
    )
    db.commit()

    return {
        "success": True,
        "score":              rapport.score,
        "alertes_mineures":   [a.message for a in rapport.alertes_mineures],
        "droits_potentiels":  rapport.droits_potentiels,
        "recommandations":    rapport.recommandations,
    }


@app.get("/api/v1/dossiers/{dossier_id}/rapport-comite")
def get_rapport_comite(
    dossier_id: str,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """Retourne le dernier rapport comité depuis cerfa_audit_log. Lecture seule."""
    row = db.execute(
        """SELECT details_json, event_at FROM cerfa_audit_log
           WHERE dossier_id = ? AND event_type = 'analyse_agents'
           ORDER BY event_at DESC LIMIT 1""",
        (dossier_id,)
    ).fetchone()
    if not row:
        return {"rapport": None, "message": "Aucun rapport disponible pour ce dossier."}
    return {"rapport": json.loads(row["details_json"] or "{}"), "date": row["event_at"]}


def _ajouter_page_couverture(cerfa_bytes: bytes, donnees: dict) -> bytes:
    """
    Génère une page de couverture avec reportlab et la fusionne AVANT le CERFA avec pypdf.
    Le CERFA officiel et ses champs ne sont pas modifiés.
    En cas d'erreur, lève une exception — l'appelant décide du fallback.
    """
    import io
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from pypdf import PdfReader, PdfWriter

    # ── Génération de la page de couverture avec reportlab ───────────────────
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        topMargin=3 * cm,
        bottomMargin=2.5 * cm,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
    )
    styles = getSampleStyleSheet()

    style_titre = ParagraphStyle(
        "titre",
        parent=styles["Heading1"],
        fontSize=14,
        textColor=colors.HexColor("#1a3a6b"),
        spaceAfter=0.3 * cm,
        alignment=TA_CENTER,
        fontName="Helvetica-Bold",
    )
    style_sous_titre = ParagraphStyle(
        "sous_titre",
        parent=styles["Normal"],
        fontSize=11,
        textColor=colors.HexColor("#1a3a6b"),
        spaceAfter=0.8 * cm,
        alignment=TA_CENTER,
        fontName="Helvetica",
    )
    style_corps = ParagraphStyle(
        "corps",
        parent=styles["Normal"],
        fontSize=10,
        leading=16,
        spaceAfter=0.5 * cm,
        alignment=TA_JUSTIFY,
        fontName="Helvetica",
    )
    style_important = ParagraphStyle(
        "important",
        parent=styles["Normal"],
        fontSize=10,
        leading=16,
        spaceAfter=0.5 * cm,
        alignment=TA_JUSTIFY,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#1a3a6b"),
    )
    style_pied = ParagraphStyle(
        "pied",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.HexColor("#666666"),
        alignment=TA_CENTER,
        fontName="Helvetica",
    )

    # Informations de contexte
    nom_prenom = donnees.get("nom_prenom", "").strip()
    reference  = donnees.get("reference", "")
    dept       = donnees.get("departement", "")
    mention_ref = f"Dossier : {reference} — Département {dept}" if reference else ""

    elements = [
        Spacer(1, 1.5 * cm),
        Paragraph("FACILIM", style_titre),
        Paragraph("Aide à la préparation du dossier MDPH", style_sous_titre),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor("#1a3a6b")),
        Spacer(1, 1 * cm),
        Paragraph("INFORMATIONS IMPORTANTES", style_important),
        Spacer(1, 0.5 * cm),
        Paragraph(
            "Ce dossier a été préparé avec l'assistance de Facilim à partir des informations "
            "déclarées par le demandeur et des documents transmis.",
            style_corps,
        ),
        Paragraph(
            "Le contenu du dossier doit être <b>relu, vérifié et validé</b> par le demandeur "
            "ou son représentant légal avant signature et transmission à la MDPH.",
            style_corps,
        ),
        Paragraph(
            "<b>La responsabilité de l'exactitude des informations appartient au signataire "
            "du dossier.</b>",
            style_important,
        ),
        Spacer(1, 0.5 * cm),
        HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc")),
        Spacer(1, 0.5 * cm),
        Paragraph(
            "Facilim est un outil numérique d'aide à la constitution de dossiers MDPH. "
            "Il ne se substitue pas à l'accompagnement professionnel ni à la responsabilité "
            "du demandeur ou de son représentant légal.",
            style_corps,
        ),
    ]

    if nom_prenom:
        elements.insert(3, Paragraph(f"Dossier préparé pour : <b>{nom_prenom}</b>", style_sous_titre))
    if mention_ref:
        elements.append(Spacer(1, 0.5 * cm))
        elements.append(Paragraph(mention_ref, style_pied))

    doc.build(elements)
    couverture_bytes = buf.getvalue()

    # ── Fusion : couverture AVANT le CERFA ───────────────────────────────────
    # CRITIQUE : on CLONE le CERFA rempli (clone_from) pour PRÉSERVER son AcroForm
    # (612 champs) et toutes les valeurs saisies, puis on insère la couverture en
    # tête. L'ancienne méthode (PdfWriter() + add_page) recopiait seulement les
    # PAGES et perdait l'AcroForm document-level → tous les champs remplis
    # disparaissaient et le CERFA sortait VIDE. (Validé par reproduction : 612→0
    # avec add_page, 612 conservés avec clone_from + insert_page.)
    writer = PdfWriter(clone_from=io.BytesIO(cerfa_bytes))

    couverture_reader = PdfReader(io.BytesIO(couverture_bytes))
    for i, page in enumerate(couverture_reader.pages):
        writer.insert_page(page, i)

    # Forcer les lecteurs PDF à afficher les valeurs des champs après fusion.
    try:
        writer.set_need_appearances_writer(True)
    except Exception:
        pass

    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def _verrou_export_cerfa(dossier_id: str, db, user, action: str = "export",
                         exiger_present: bool = False) -> dict:
    """
    VERROU CENTRALISÉ d'export CERFA (FACILIM PROD-3).

    Lit le gate DÉJÀ calculé dans `dossiers.cockpit_pro_json` (aucun recalcul,
    aucune nouvelle évaluation, aucun moteur appelé), journalise la décision dans
    `cerfa_audit_log`, et lève HTTP 403 si le gate est en BLOCK.

      - PASS / WARNING / UNKNOWN(absent) → autorise (retourne le statut)
      - BLOCK                            → refuse (403, aucun PDF généré)

    exiger_present : durcissement réservé à la TRANSMISSION officielle. Quand True,
    un gate absent (cockpit jamais évalué) ne suffit plus : le dossier doit avoir
    été évalué et validé avant d'être transmis. Le téléchargement pour relecture
    reste, lui, tolérant (fail-open) afin de ne pas gêner le travail du pro.
    """
    from app.services.export_gate import lire_gate_export, coeur_incomplet

    row = db.execute(
        "SELECT cockpit_pro_json, synthese_json FROM dossiers WHERE id = ?", (dossier_id,)
    ).fetchone()
    cockpit_json = row["cockpit_pro_json"] if row else None
    g = lire_gate_export(cockpit_json)

    # SÉCURISATION FRAÎCHEUR (anti-gate-obsolète) — le gate ci-dessus est PERSISTÉ et
    # peut être périmé (une édition dashboard modifie synthese_json sans recalculer le
    # gate). On revérifie donc la solidité du CŒUR sur la synthèse FRAÎCHE, à chaque
    # export. Déterministe, sans LLM, sans toucher aux moteurs droits/éligibilité.
    try:
        _synthese_fraiche = json.loads((row["synthese_json"] if row else "{}") or "{}")
    except Exception:
        _synthese_fraiche = {}
    _coeur_bloque, _coeur_raison = coeur_incomplet(_synthese_fraiche)

    acteur = None
    try:
        acteur = (user or {}).get("sub")
    except Exception:
        acteur = None

    _log_cerfa_audit(
        db, dossier_id,
        event_type="gate_export_check",
        canal="dashboard",
        acteur_id=acteur,
        acteur_type="professionnel",
        details={
            "action":          action,
            "gate_statut":     g["statut"],
            "gate_present":    g["present"],
            "export_autorise": g["autorise_export"],
            "raison":          g["message"] if not g["autorise_export"] else "",
        },
    )
    try:
        db.commit()
    except Exception:
        pass

    # Blocage FRAÎCHEUR : un cœur métier devenu incomplet bloque l'export même si le
    # gate persisté est encore « PASS » (anti-désynchronisation). Indépendant du gate.
    if _coeur_bloque:
        raise HTTPException(
            status_code=403,
            detail=("Export bloqué : " + _coeur_raison + ". Le cœur du dossier (retentissement, "
                    "aides, attentes, projet de vie) doit être réellement renseigné avant export "
                    "ou transmission. Complétez le dossier puis réessayez."),
        )

    if not g["autorise_export"]:
        raise HTTPException(
            status_code=403,
            detail=(g["message"] or
                    "Export bloqué par FACILIM (gate BLOCK). Corrigez les points critiques "
                    "signalés dans le cockpit avant d'exporter ou de transmettre le CERFA."),
        )
    # Durcissement transmission : un dossier jamais évalué ne peut pas être transmis.
    if exiger_present and not g["present"]:
        raise HTTPException(
            status_code=403,
            detail=("Le dossier n'a pas encore été évalué par FACILIM. Ouvrez-le dans le "
                    "cockpit pour générer l'analyse, faites valider par un professionnel, "
                    "puis relancez la transmission."),
        )
    return g


def _generer_cerfa_pdf(dossier_id: str, db) -> tuple[str, bytes]:
    """
    Génère le PDF CERFA pour un dossier et retourne (pdf_path, pdf_bytes).

    UNIQUE moteur de remplissage : V2 (services/cerfa_filler.py), via le pont
    v2_bridge.py. Il n'y a plus de second moteur de secours : en cas d'échec V2,
    on remonte une erreur claire plutôt que de produire silencieusement un CERFA
    issu d'un autre moteur, source d'incohérences entre dossiers.
    """
    import os
    from datetime import datetime, timezone

    dossier_row = db.execute("SELECT * FROM dossiers WHERE id = ?", (dossier_id,)).fetchone()
    if not dossier_row:
        raise HTTPException(status_code=404, detail="Dossier non trouvé")

    dossier  = dict(dossier_row)
    donnees  = json.loads(dossier.get("synthese_json", "{}") or "{}")
    session  = db.execute(
        "SELECT persona_actif FROM sessions_whatsapp WHERE dossier_id = ? LIMIT 1",
        (dossier_id,)
    ).fetchone()
    service_type = (session["persona_actif"] if session else "adulte") or "adulte"
    num_mdph     = donnees.get("numero_dossier_mdph", "")
    type_dos     = dossier.get("type_dossier", donnees.get("type_dossier", "INITIAL"))

    # Stocker le type de dossier dans donnees pour le bridge
    if type_dos:
        donnees["type_dossier"] = type_dos

    # Sprint P0.6 — Enrichir donnees depuis colonnes DB si absentes de synthese_json
    # Ces champs peuvent être stockés en colonne DB mais pas dans synthese_json.
    if not donnees.get("numero_dossier_mdph"):
        _ndm = dossier.get("numero_dossier_mdph", "")
        if _ndm:
            donnees["numero_dossier_mdph"] = _ndm
    if not donnees.get("type_dossier") or donnees.get("type_dossier") == "INITIAL":
        _td = dossier.get("type_dossier", "")
        if _td and _td != "INITIAL":
            donnees["type_dossier"] = _td

    # ── Contrôle qualité métier avant génération ────────────────────────────────
    _alertes_cerfa = []

    # Contrôle A1
    if not donnees.get("nom_prenom"):
        _alertes_cerfa.append("A1 — Nom/prénom manquant")
    if not (donnees.get("num_secu") or donnees.get("numero_securite_sociale")):
        _alertes_cerfa.append("A1 — NIR non renseigné (à collecter)")
    if not donnees.get("genre"):
        _alertes_cerfa.append("A1 — Genre non renseigné")

    # Contrôle B1 — contradiction aide humaine
    _dit_pas_besoin_aide = any(
        w in str(donnees.get("impact_quotidien", "")).lower()
        for w in ["pas besoin", "pas d'aide", "sans aide", "autonome", "ne nécessite pas"]
    )
    if _dit_pas_besoin_aide and donnees.get("besoins_aide_humaine"):
        _alertes_cerfa.append(
            "B2 — CONTRADICTION : usager dit ne pas avoir besoin d'aide humaine "
            "mais besoins_aide_humaine=True — validation humaine requise"
        )
        # Flag humain automatique
        from app.engines.case_state_engine import create_flag_humain
        try:
            create_flag_humain(
                dossier_id=dossier_id,
                raison="Contradiction B2 : déclaration vs aide humaine — vérification requise",
                educateur_id=None, severite="NORMALE", db_conn=db,
            )
        except Exception:
            pass

    # Contrôle D — accident travail → a déjà travaillé
    if donnees.get("accident_travail") and str(donnees.get("a_deja_travaille", "")).lower() == "false":
        donnees["a_deja_travaille"] = True
        _alertes_cerfa.append("D — Auto-correction : accident_travail détecté → a_deja_travaille=True")

    # Contrôle D — France Travail / Pôle Emploi
    _ft_mentions = ["france travail", "pôle emploi", "pole emploi"]
    if any(w in str(donnees.get("statut_emploi", "")).lower() for w in _ft_mentions):
        donnees["inscrit_pole_emploi"] = True

    if _alertes_cerfa:
        logger.info("[CERFA QC] %d alerte(s) détectée(s) pour dossier %s : %s",
                    len(_alertes_cerfa), dossier_id[:8], "; ".join(_alertes_cerfa))

    storage_path = "/data/cerfa" if os.path.isdir("/data") else "./storage/cerfa"
    os.makedirs(storage_path, exist_ok=True)

    output_filename = (
        f"cerfa_{dossier_id[:8]}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.pdf"
    )
    output_path = os.path.join(storage_path, output_filename)

    # ── Moteur V2 (unique) ──────────────────────────────────────────────────
    # L'audit (hash, SELECT version, journalisation) et la persistance sont NON
    # BLOQUANTS — une erreur d'audit (ex. « no such column: version ») ne doit
    # JAMAIS détruire un PDF V2 déjà généré.
    pdf_bytes = None
    try:
        from app.engines.pdf.v2_bridge import generer_cerfa_depuis_synthese
        pdf_bytes = generer_cerfa_depuis_synthese(donnees, service_type, num_mdph)
    except Exception as e_v2:
        logger.error("[CERFA] Génération V2 échouée (%s) — aucune bascule, erreur remontée", e_v2)
        pdf_bytes = None

    if pdf_bytes is not None:
        # ── Hashes audit (sur le CERFA officiel, AVANT couverture) — non bloquant ─
        _synthese_hash = _cerfa_pdf_hash = ""
        _nb_pages = 0
        try:
            import hashlib as _hl, json as _j_hash
            _synthese_hash  = _hl.sha256(
                _j_hash.dumps(donnees, sort_keys=True, ensure_ascii=False).encode("utf-8")
            ).hexdigest()
            _cerfa_pdf_hash = _hl.sha256(pdf_bytes).hexdigest()
            from pypdf import PdfReader as _PR
            _nb_pages = len(_PR(io.BytesIO(pdf_bytes)).pages)
        except Exception as e_hash:
            logger.debug("[CERFA V2] Hash/pagination audit ignoré : %s", e_hash)

        # ── Page de couverture (mise en forme) — non bloquant ─────────────────
        try:
            pdf_bytes = _ajouter_page_couverture(pdf_bytes, donnees)
            logger.info("[CERFA] Page de couverture ajoutée | dossier=%s", dossier_id[:8])
        except Exception as e_cov:
            logger.warning("[CERFA] Page de couverture ignorée (%s) — PDF CERFA seul transmis", e_cov)

        # ── Journalisation audit (Point 3) — NON BLOQUANTE ────────────────────
        # Une erreur SQL / audit / version ici ne doit JAMAIS détruire le PDF V2.
        try:
            _dossier_version = db.execute(
                "SELECT version FROM dossiers WHERE id = ?", (dossier_id,)
            ).fetchone()
            _version_val = (_dossier_version["version"] if _dossier_version and _dossier_version["version"] else 1)
            _log_cerfa_audit(
                db=db,
                dossier_id=dossier_id,
                event_type="generation_cerfa",
                canal="dashboard",
                acteur_id=None,
                acteur_type="systeme",
                details={
                    "synthese_hash": _synthese_hash,
                    "pdf_hash":      _cerfa_pdf_hash,
                    "version":       _version_val,
                    "nb_pages":      _nb_pages,
                },
            )
        except Exception as e_audit:
            logger.warning("[CERFA V2] Audit non bloquant ignoré (%s) — PDF V2 conservé", e_audit)

        # ── Persistance du PDF V2 ─────────────────────────────────────────────
        with open(output_path, "wb") as f:
            f.write(pdf_bytes)
        logger.info("[CERFA V2] PDF généré via moteur V2 | dossier=%s | %d bytes | hash=%s…",
                    dossier_id[:8], len(pdf_bytes), (_synthese_hash[:8] or "—"))
        return output_path, pdf_bytes

    # ── Échec V2 : moteur UNIQUE, donc aucune bascule silencieuse vers un autre
    #    moteur. On remonte une erreur explicite : mieux vaut pas de document qu'un
    #    CERFA produit par un moteur divergent, source d'incohérences.
    logger.error(
        "[CERFA] Génération indisponible pour dossier=%s — aucune bascule de moteur, "
        "erreur remontée au professionnel.", dossier_id[:8],
    )
    raise HTTPException(
        status_code=503,
        detail=("La génération du CERFA a échoué temporairement. Vérifiez les données du "
                "dossier puis réessayez ; aucun document partiel n'est produit."),
    )


@app.get("/api/v1/dossiers/{dossier_id}/cerfa.pdf")
def telecharger_cerfa(
    dossier_id: str,
    export: bool = False,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """
    Sert le PDF CERFA.

    - `export=false` (défaut) → PRÉVISUALISATION / consultation : TOUJOURS disponible,
      jamais bloquée par le gate (le PDF s'ouvre inline dans le navigateur).
    - `export=true` → EXPORT DÉFINITIF : passe par le verrou PROD-3 ; refusé (403)
      si le gate est en BLOCK (aucun PDF généré).
    """
    from fastapi.responses import Response as RawResponse
    try:
        # Verrou uniquement pour l'export définitif — la prévisualisation reste libre.
        if export:
            _verrou_export_cerfa(dossier_id, db, user, action="telechargement")
        _, pdf_bytes = _generer_cerfa_pdf(dossier_id, db)
        nom_usager  = "dossier"
        row = db.execute("SELECT synthese_json FROM dossiers WHERE id = ?", (dossier_id,)).fetchone()
        if row:
            d = json.loads(row["synthese_json"] or "{}")
            import unicodedata as _ud_hdr
            _nom_brut = str(d.get("nom_prenom", "") or "dossier")
            # L'en-tête HTTP Content-Disposition n'accepte que le latin-1. Les
            # apostrophes courbes (') et autres caractères Unicode font planter
            # l'envoi. On translittère donc le nom en ASCII pur pour le nom de fichier.
            _nom_ascii = _ud_hdr.normalize("NFKD", _nom_brut).encode("ascii", "ignore").decode("ascii")
            nom_usager = "".join(
                c if (c.isalnum() or c in "-_") else "_" for c in _nom_ascii
            ).strip("_") or "dossier"
        return RawResponse(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'inline; filename="CERFA_{nom_usager}.pdf"'},
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error("[CERFA] Erreur prévisualisation : %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/v1/dossiers/{dossier_id}/previsualiser-cerfa")
def previsualiser_cerfa(
    dossier_id: str,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """Retourne l'URL de téléchargement du CERFA pour ouverture dans le navigateur."""
    # On vérifie juste que le dossier existe — la génération se fait via GET /cerfa.pdf
    row = db.execute("SELECT id FROM dossiers WHERE id = ?", (dossier_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Dossier non trouvé")
    return {
        "success":   True,
        "pdf_url":   f"/api/v1/dossiers/{dossier_id}/cerfa.pdf",
        "message":   "Cliquez sur pdf_url pour ouvrir le CERFA dans votre navigateur.",
    }


@app.post("/api/v1/dossiers/{dossier_id}/envoyer-cerfa")
def envoyer_cerfa(
    dossier_id: str,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """
    Génère le CERFA PDF et l'envoie par email à l'usager.
    L'adresse email est lue depuis synthese_json (renseignée à la création du dossier).
    """
    import requests as _req

    # Récupérer l'email de l'usager
    row = db.execute("SELECT synthese_json, reference FROM dossiers WHERE id = ?", (dossier_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Dossier non trouvé")

    # ── VERROU PROD-3 : gate d'export (lit le gate DÉJÀ calculé, aucun recalcul) ──
    # BLOCK → 403 ici : aucun PDF n'est généré, aucun email n'est envoyé.
    gate = _verrou_export_cerfa(dossier_id, db, user, action="envoi", exiger_present=True)

    donnees    = json.loads(row["synthese_json"] or "{}")
    email_dest = donnees.get("email", "").strip()
    reference  = row["reference"] or dossier_id[:8]

    if not email_dest:
        raise HTTPException(
            status_code=400,
            detail="Aucun email trouvé pour cet usager. Ajoutez l'email dans le dossier avant d'envoyer.",
        )

    # ── B4 : Validation usager requise avant envoi ────────────────────────────
    _val_usager = db.execute(
        """SELECT id FROM cerfa_validations
           WHERE dossier_id = ? AND type_validation = 'usager' AND reponse_usager = 'OUI'
           ORDER BY validated_at DESC LIMIT 1""",
        (dossier_id,)
    ).fetchone()
    if not _val_usager:
        raise HTTPException(
            status_code=400,
            detail=(
                "Validation usager requise avant envoi. "
                "Le bénéficiaire doit confirmer la relecture du dossier via WhatsApp."
            ),
        )

    # ── B5 : Validation professionnelle requise avant envoi ───────────────────
    _val_pro = db.execute(
        """SELECT id FROM cerfa_validations
           WHERE dossier_id = ? AND type_validation = 'professionnel' AND reponse_usager = 'OUI'
           ORDER BY validated_at DESC LIMIT 1""",
        (dossier_id,)
    ).fetchone()
    if not _val_pro:
        raise HTTPException(
            status_code=400,
            detail=(
                "Validation professionnelle requise avant envoi. "
                "Utilisez le bouton 'Valider la remise' dans le dossier."
            ),
        )

    # Générer le PDF
    try:
        _, pdf_bytes = _generer_cerfa_pdf(dossier_id, db)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Génération PDF échouée : {e}")

    # Envoi via Brevo HTTP API (avec pièce jointe)
    import base64
    pdf_b64   = base64.b64encode(pdf_bytes).decode()
    nom_dest  = donnees.get("nom_prenom", "").split()[0] if donnees.get("nom_prenom") else ""
    civilite  = f"Bonjour{' ' + nom_dest if nom_dest else ''}"

    payload = {
        "sender":      {"name": settings.brevo_sender_name, "email": settings.brevo_sender_email},
        "to":          [{"email": email_dest}],
        "subject":     f"[Facilim] Votre formulaire MDPH — Dossier {reference}",
        "htmlContent": f"""
            <p>{civilite},</p>
            <p>Vous trouverez ci-joint votre dossier MDPH préparé à partir des informations
            déclarées et des documents transmis dans le cadre de votre accompagnement.</p>
            <p>Nous vous invitons à <strong>relire attentivement l'ensemble du dossier</strong>
            avant toute signature ou transmission à la MDPH afin de vérifier l'exactitude,
            l'exhaustivité et l'actualité des informations figurant dans le document.</p>
            <p>Facilim est un outil d'aide à la préparation du dossier.<br/>
            La validation finale des informations et la signature du CERFA relèvent
            exclusivement du demandeur ou de son représentant légal.</p>
            <p>En cas d'erreur, d'information manquante ou de modification à apporter,
            merci de nous en informer avant toute transmission à la MDPH.</p>
            <p>Cordialement,<br/>L'équipe Facilim</p>
        """,
        "attachment":  [{"content": pdf_b64, "name": f"CERFA_MDPH_{reference}.pdf"}],
    }
    resp = _req.post(
        "https://api.brevo.com/v3/smtp/email",
        headers={"api-key": settings.brevo_api_key, "content-type": "application/json"},
        json=payload,
        timeout=15,
    )
    if resp.status_code not in (200, 201):
        logger.error("[CERFA] Envoi email échoué : %s %s", resp.status_code, resp.text[:200])
        raise HTTPException(status_code=500, detail=f"Envoi email échoué ({resp.status_code})")

    # Mettre à jour le statut
    now = datetime.now(timezone.utc).isoformat()
    db.execute(
        "UPDATE dossiers SET statut = 'SOUMIS', soumis_le = ?, updated_at = ? WHERE id = ?",
        (now, now, dossier_id),
    )
    event_logger.log_event("CERFA_ENVOYE_USAGER", dossier_id=dossier_id,
                           utilisateur_id=user.get("sub"), canal="dashboard",
                           payload={"email": email_dest[-10:]}, db_conn=db)
    return {
        "success": True,
        "email_dest": email_dest,
        "soumis_le": now,
        "gate_statut": gate["statut"],
        "gate_avertissement": gate["message"] if gate["statut"] == "WARNING" else "",
    }


class CerfaChampRequest(BaseModel):
    champ:  str
    valeur: str = ""


@app.post("/api/v1/dossiers/{dossier_id}/synthese-update")
def bulk_update_synthese(
    dossier_id: str,
    body: dict,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """
    Met à jour plusieurs champs de synthese_json en une seule requête.
    Les champs existants ne sont écrasés QUE si la nouvelle valeur est non vide.
    Utilisé par : saisie manuelle dashboard, retour email, corrections pro.
    """
    dossier_row = db.execute("SELECT synthese_json FROM dossiers WHERE id = ?", (dossier_id,)).fetchone()
    if not dossier_row:
        raise HTTPException(status_code=404, detail="Dossier non trouvé")

    synthese = json.loads(dossier_row["synthese_json"] or "{}")
    champs_mis_a_jour = []

    for champ, valeur in body.items():
        if champ.startswith("_"):  # ignorer les champs internes
            continue
        if valeur is not None and valeur != "":
            if champ == "notes_pro" and synthese.get("notes_pro"):
                synthese["notes_pro"] = f"{synthese['notes_pro']}\n{valeur}".strip()
            else:
                synthese[champ] = valeur
            champs_mis_a_jour.append(champ)

    # Alias NIR : garder les noms du numéro de sécu synchronisés (l'interface peut
    # enregistrer l'un ou l'autre ; la collecte WhatsApp lit 'num_secu'). Helper partagé.
    from app.services.collecte_schema import synchroniser_nir
    synchroniser_nir(synthese)

    db.execute(
        "UPDATE dossiers SET synthese_json = ?, updated_at = ? WHERE id = ?",
        (json.dumps(synthese, ensure_ascii=False), datetime.now(timezone.utc).isoformat(), dossier_id),
    )
    logger.info("[SYNTHESE] %d champs mis à jour pour dossier %s : %s",
                len(champs_mis_a_jour), dossier_id, champs_mis_a_jour)
    return {"success": True, "champs_mis_a_jour": champs_mis_a_jour}


@app.post("/api/v1/dossiers/{dossier_id}/cerfa-champ")
def update_cerfa_champ(
    dossier_id: str,
    body: CerfaChampRequest,          # Pydantic — plus besoin de asyncio.get_event_loop()
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """Met à jour un champ ou une note pro dans le dossier depuis le dashboard."""
    if not body.champ:
        raise HTTPException(status_code=400, detail="Champ manquant")
    dossier_row = db.execute("SELECT synthese_json FROM dossiers WHERE id = ?", (dossier_id,)).fetchone()
    if not dossier_row:
        raise HTTPException(status_code=404, detail="Dossier non trouvé")

    synthese = json.loads(dossier_row["synthese_json"] or "{}")

    # Si c'est une note pro → concaténer à la valeur existante
    if body.champ == "notes_pro" and synthese.get("notes_pro"):
        synthese["notes_pro"] = f"{synthese['notes_pro']}\n{body.valeur}".strip()
    else:
        synthese[body.champ] = body.valeur

    # Alias NIR : l'interface enregistre le numéro de Sécu sous 'numero_securite_sociale',
    # mais la collecte WhatsApp surveille 'num_secu'. On synchronise les noms pour qu'un
    # NIR saisi à l'écran ne soit JAMAIS redemandé sur WhatsApp. Helper partagé (gère 'nss').
    from app.services.collecte_schema import synchroniser_nir
    synchroniser_nir(synthese)

    from app.engines.orchestration_engine import _score_completude_metier
    score = _score_completude_metier(synthese)   # P4 : complétude métier (synthese_completude)
    db.execute(
        """UPDATE dossiers SET synthese_json = ?,
           score_completude = CASE WHEN score_completude < ? THEN ? ELSE score_completude END,
           updated_at = ? WHERE id = ?""",
        (json.dumps(synthese, ensure_ascii=False), score, score,
         datetime.now(timezone.utc).isoformat(), dossier_id),
    )
    return {"success": True, "champ": body.champ}


@app.get("/api/v1/dossiers/{dossier_id}/score")
def get_score(
    dossier_id: str,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """Retourne le dernier scoring calculé pour un dossier."""
    scoring = db.execute(
        "SELECT * FROM analyses_scoring WHERE dossier_id = ? ORDER BY created_at DESC LIMIT 1",
        (dossier_id,),
    ).fetchone()
    if not scoring:
        return {"score_global": 0, "confiance": 0.0, "statut_analyse": "NON_CALCULE",
                "droits_identifies": [], "recommandation": None}
    s = dict(scoring)
    s["droits_identifies"] = json.loads(s.get("droits_identifies_json", "[]") or "[]")
    return s


@app.post("/api/v1/dossiers/{dossier_id}/email-medical-recu")
def email_medical_recu(
    dossier_id: str,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """Confirme la réception du certificat médical."""
    now = datetime.now(timezone.utc).isoformat()
    db.execute(
        "UPDATE dossiers SET certificat_medical_present = 1, updated_at = ? WHERE id = ?",
        (now, dossier_id),
    )
    event_logger.log_event("CERTIFICAT_MEDICAL_RECU", dossier_id=dossier_id,
                           utilisateur_id=user.get("sub"), canal="dashboard", db_conn=db)
    return {"success": True}


@app.delete("/api/v1/dossiers/{dossier_id}")
def delete_dossier(
    dossier_id: str,
    user=Depends(_get_current_user),
    db=Depends(_get_db),
):
    """Soft-delete d'un dossier."""
    access_control.require_permission(user.get("role"), access_control.Permission.HUMAN_VALIDATE)
    now = datetime.now(timezone.utc).isoformat()
    db.execute("UPDATE dossiers SET deleted_at = ?, updated_at = ? WHERE id = ?",
               (now, now, dossier_id))
    event_logger.log_event("DOSSIER_SUPPRIME", dossier_id=dossier_id,
                           utilisateur_id=user.get("sub"), canal="dashboard", db_conn=db)
    return {"success": True}
